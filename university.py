import os
import json
import time
import re
import threading
import hashlib
import tkinter as tk
from tkinter import ttk, messagebox, scrolledtext, font
from functools import wraps
from typing import List, Optional, Dict, Any
from dataclasses import dataclass
from pydantic import BaseModel
import logging
from datetime import datetime
import webbrowser
import subprocess
import platform

# Import config
try:
    import config
except ImportError:
    print("Please create config.py with your GEMINI_API_KEY")
    exit(1)

# Setup logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s [%(levelname)s] %(message)s')
logger = logging.getLogger(__name__)

# Configure Gemini
import google.generativeai as genai
genai.configure(api_key=getattr(config, 'GEMINI_API_KEY', ''))

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_community.vectorstores import Chroma
from langchain.schema import Document

import PyPDF2
from docx import Document as DocxDocument

# Color scheme - University theme
COLORS = {
    'primary': '#1e3a8a',
    'secondary': '#7c3aed',
    'success': '#059669',
    'danger': '#dc2626',
    'warning': '#d97706',
    'info': '#0891b2',
    'light': '#f8fafc',
    'dark': '#1e293b',
    'white': '#ffffff',
    'bg_primary': '#dbeafe',
    'bg_secondary': '#e0e7ff',
    'text_dark': '#1e293b',
    'text_light': '#64748b',
    'border': '#e2e8f0'
}

# Configuration
UNIVERSITY_DOCS_DIR = "./university_documents"
VECTOR_STORE_DIR = "./university_vector_store"
API_RATE_LIMIT = getattr(config, 'API_RATE_LIMIT', 10)
MAX_RETRIES = getattr(config, 'MAX_RETRIES', 3)
RETRY_BASE_DELAY = getattr(config, 'RETRY_BASE_DELAY', 5)
CHUNK_SIZE = getattr(config, 'CHUNK_SIZE', 1000)
CHUNK_OVERLAP = getattr(config, 'CHUNK_OVERLAP', 200)
MAX_CONTEXT_DOCS = getattr(config, 'MAX_CONTEXT_DOCS', 5)
CACHE_DIR = getattr(config, 'CACHE_DIR', "./api_cache")

# ------------------ Core System Classes (Keep same as before) ------------------
def rate_limit(calls_per_minute: int):
    min_interval = 60.0 / calls_per_minute
    
    def decorator(func):
        last_called = [0.0]
        
        @wraps(func)
        def wrapper(*args, **kwargs):
            elapsed = time.time() - last_called[0]
            wait_time = min_interval - elapsed
            if wait_time > 0:
                time.sleep(wait_time)
            ret = func(*args, **kwargs)
            last_called[0] = time.time()
            return ret
        return wrapper
    return decorator

class APICache:
    def __init__(self, cache_dir=CACHE_DIR):
        self.cache_dir = cache_dir
        os.makedirs(self.cache_dir, exist_ok=True)
    
    def _get_cache_key(self, prompt: str) -> str:
        return hashlib.md5(prompt.encode('utf-8')).hexdigest()
    
    def get(self, prompt: str) -> Optional[str]:
        cache_key = self._get_cache_key(prompt)
        cache_file = os.path.join(self.cache_dir, f"{cache_key}.json")
        
        if os.path.exists(cache_file):
            try:
                with open(cache_file, "r", encoding="utf-8") as f:
                    data = json.load(f)
                return data.get("content")
            except Exception:
                pass
        return None
    
    def set(self, prompt: str, content: str):
        cache_key = self._get_cache_key(prompt)
        cache_file = os.path.join(self.cache_dir, f"{cache_key}.json")
        
        try:
            with open(cache_file, "w", encoding="utf-8") as f:
                json.dump({"content": content, "timestamp": time.time()}, f)
        except Exception:
            pass

@dataclass
class QueryStructure:
    department: Optional[str] = None
    program: Optional[str] = None
    semester: Optional[str] = None
    course: Optional[str] = None
    topic: Optional[str] = None
    raw_query: str = ""

class AnalysisResponse(BaseModel):
    answer: str
    detailed_answer: str
    justification: str
    confidence_score: float
    key_points: List[str]
    document_references: List[str]
    sources: List[Dict[str, Any]]
    applicable_sections: List[str]

class SafeModelClient:
    def __init__(self, model_name: str = "gemini-1.5-flash"):
        self.model = genai.GenerativeModel(model_name)
        self.cache = APICache()
        self.rate_limited_generate = rate_limit(API_RATE_LIMIT)(self._generate_content_internal)
    
    def _generate_content_internal(self, prompt: str) -> str:
        response = self.model.generate_content(prompt)
        return response.text or ""
    
    def generate_content(self, prompt: str) -> str:
        cached_content = self.cache.get(prompt)
        if cached_content:
            return cached_content
        
        for attempt in range(MAX_RETRIES):
            try:
                content = self.rate_limited_generate(prompt)
                self.cache.set(prompt, content)
                return content
            except Exception as e:
                if attempt == MAX_RETRIES - 1:
                    raise e
                time.sleep(RETRY_BASE_DELAY * (2 ** attempt))
        
        raise RuntimeError("Exceeded maximum retry attempts")

class UniversityDocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE,
            chunk_overlap=CHUNK_OVERLAP,
            separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""]
        )
    
    def load_pdf(self, filepath: str) -> str:
        try:
            text = ""
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text += page_text + "\n"
                    except Exception:
                        continue
            return text.strip() if text.strip() else None
        except Exception:
            return None
    
    def load_docx(self, filepath: str) -> str:
        try:
            doc = DocxDocument(filepath)
            text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text.strip() + "\n"
            return text.strip() if text.strip() else None
        except Exception:
            return None
    
    def load_text(self, filepath: str) -> str:
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                return f.read().strip()
        except Exception:
            return None
    
    def load_all_university_documents(self) -> List[Document]:
        documents = []
        
        if not os.path.exists(UNIVERSITY_DOCS_DIR):
            return documents
        
        files = [f for f in os.listdir(UNIVERSITY_DOCS_DIR) if not os.path.isdir(os.path.join(UNIVERSITY_DOCS_DIR, f))]
        
        for filename in files:
            filepath = os.path.join(UNIVERSITY_DOCS_DIR, filename)
            file_ext = os.path.splitext(filepath)[1].lower()
            
            text = None
            if file_ext == '.pdf':
                text = self.load_pdf(filepath)
            elif file_ext == '.docx':
                text = self.load_docx(filepath)
            elif file_ext in ['.txt', '.md']:
                text = self.load_text(filepath)
            
            if not text:
                continue
            
            # Clean text
            text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
            text = re.sub(r' +', ' ', text)
            text = text.strip()
            
            # Split into chunks
            chunks = self.text_splitter.split_text(text)
            
            for i, chunk in enumerate(chunks):
                if chunk.strip():
                    doc = Document(
                        page_content=chunk,
                        metadata={
                            "source": filepath,
                            "filename": filename,
                            "chunk_id": f"{filename}_{i}",
                            "file_type": file_ext
                        }
                    )
                    documents.append(doc)
        
        return documents

class UniversityQueryProcessor:
    def __init__(self):
        self.model_client = SafeModelClient('gemini-1.5-flash')
        self.embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
        self.vector_store = None
        self.system_initialized = False
    
    def initialize_system(self, documents: List[Document]):
        try:
            # Try to load existing vector store
            if os.path.exists(VECTOR_STORE_DIR) and os.listdir(VECTOR_STORE_DIR):
                self.vector_store = Chroma(
                    persist_directory=VECTOR_STORE_DIR,
                    embedding_function=self.embeddings
                )
                self.system_initialized = True
                return
            
            # Create new vector store
            if documents:
                self.vector_store = Chroma.from_documents(
                    documents=documents,
                    embedding=self.embeddings,
                    persist_directory=VECTOR_STORE_DIR
                )
                self.system_initialized = True
            
        except Exception as e:
            raise Exception(f"Error initializing system: {e}")
    
    def search_relevant_content(self, query: str, k: int = 5) -> List[Document]:
        if not self.vector_store:
            return []
        return self.vector_store.similarity_search(query, k=k)
    
    def process_query(self, query: str) -> AnalysisResponse:
        if not self.system_initialized:
            raise ValueError("System not ready")
        
        relevant_docs = self.search_relevant_content(query, k=5)
        
        if not relevant_docs:
            return AnalysisResponse(
                answer="No relevant information found.",
                detailed_answer="No relevant information was found in the university documents for your query. Please try rephrasing your question or asking about a different topic.",
                justification="Could not find relevant content.",
                confidence_score=0.0,
                key_points=[],
                document_references=[],
                sources=[],
                applicable_sections=[]
            )
        
        # Prepare context
        context = "\n\n".join([
            f"Document: {doc.metadata.get('filename', 'Unknown')}\nContent: {doc.page_content}"
            for doc in relevant_docs[:3]
        ])
        
        # Generate structured answer
        answer_prompt = f"""
        Based on the following university documents, provide a comprehensive answer to the user's question.
        
        QUESTION: {query}
        
        CONTEXT: {context}
        
        Please provide your answer in a structured format with:
        1. Key information highlighted with **bold** formatting
        2. Clear sections and bullet points where appropriate
        3. Specific references to the source documents
        
        Make the answer well-organized and easy to read.
        """
        
        # Generate detailed explanation
        detailed_prompt = f"""
        Based on the following university documents, provide a detailed, comprehensive explanation for the user's question.
        
        QUESTION: {query}
        
        CONTEXT: {context}
        
        Please provide:
        1. **In-depth analysis** of the topic with background information
        2. **Step-by-step explanations** where applicable
        3. **Additional related information** that might be helpful
        4. **Specific examples** from the documents
        5. **Cross-references** to related topics or procedures
        6. **Practical advice** or next steps for the user
        
        This should be a thorough explanation that goes beyond the basic answer and provides comprehensive understanding of the topic.
        Format with **bold** keywords for better readability.
        """
        
        try:
            # Generate both answers
            response_text = self.model_client.generate_content(answer_prompt)
            detailed_response = self.model_client.generate_content(detailed_prompt)
            
            sources = []
            for i, doc in enumerate(relevant_docs):
                # Convert relative path to absolute path
                absolute_path = os.path.abspath(doc.metadata.get('source', ''))
                sources.append({
                    'source_id': f"source_{i+1}",
                    'filename': doc.metadata.get('filename', 'unknown'),
                    'filepath': absolute_path,
                    'content_preview': doc.page_content[:200] + "...",
                    'relevance': 1.0 - (i * 0.2)
                })
            
            return AnalysisResponse(
                answer=response_text,
                detailed_answer=detailed_response,
                justification=f"Answer based on {len(relevant_docs)} relevant documents",
                confidence_score=0.8,
                key_points=[response_text[:100] + "..."],
                document_references=[doc.metadata.get('filename', 'Unknown') for doc in relevant_docs],
                sources=sources,
                applicable_sections=[]
            )
            
        except Exception as e:
            return AnalysisResponse(
                answer=f"Error processing query: {str(e)}",
                detailed_answer=f"Error processing query: {str(e)}",
                justification="An error occurred",
                confidence_score=0.0,
                key_points=[],
                document_references=[],
                sources=[],
                applicable_sections=[]
            )

# ------------------ COMPLETELY REDESIGNED GUI WITH SCROLLABLE INTERFACE ------------------
class ScrollableFrame(tk.Frame):
    """A scrollable frame that contains a canvas and scrollbar"""
    def __init__(self, parent):
        super().__init__(parent)
        
        # Create canvas and scrollbar
        self.canvas = tk.Canvas(self, bg=COLORS['light'], highlightthickness=0)
        self.scrollbar = ttk.Scrollbar(self, orient="vertical", command=self.canvas.yview)
        self.scrollable_frame = tk.Frame(self.canvas, bg=COLORS['light'])
        
        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all"))
        )
        
        self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        self.canvas.configure(yscrollcommand=self.scrollbar.set)
        
        # Pack canvas and scrollbar
        self.canvas.pack(side="left", fill="both", expand=True)
        self.scrollbar.pack(side="right", fill="y")
        
        # Bind mousewheel to canvas
        self.bind_mousewheel()
        
        # Configure canvas width to match window
        self.canvas.bind('<Configure>', self._on_canvas_configure)
    
    def _on_canvas_configure(self, event):
        canvas_width = event.width
        self.canvas.itemconfig(self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw"), width=canvas_width)
    
    def bind_mousewheel(self):
        def _on_mousewheel(event):
            self.canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        
        def _bind_to_mousewheel(event):
            self.canvas.bind_all("<MouseWheel>", _on_mousewheel)
        
        def _unbind_from_mousewheel(event):
            self.canvas.unbind_all("<MouseWheel>")
        
        self.canvas.bind('<Enter>', _bind_to_mousewheel)
        self.canvas.bind('<Leave>', _unbind_from_mousewheel)

class EnhancedUniversityGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("🎓 Enhanced University Information Assistant")
        self.root.geometry("1600x1000")
        self.root.configure(bg=COLORS['light'])
        
        # Initialize variables
        self.status_var = tk.StringVar()
        self.status_var.set("Starting application...")
        
        # Initialize components
        self.document_processor = UniversityDocumentProcessor()
        self.query_processor = UniversityQueryProcessor()
        self.system_ready = False
        self.current_result = None
        self.detailed_visible = False
        self.current_detailed_answer = ""
        self.document_paths = {}
        
        # Create custom fonts
        self.setup_fonts()
        
        # Create GUI with improved layout
        self.create_widgets()
        
        # Initialize system
        self.initialize_system()
    
    def setup_fonts(self):
        """Setup custom fonts for better text formatting"""
        self.title_font = font.Font(family="Segoe UI", size=18, weight="bold")
        self.subtitle_font = font.Font(family="Segoe UI", size=14, weight="bold")
        self.normal_font = font.Font(family="Segoe UI", size=12)
        self.bold_font = font.Font(family="Segoe UI", size=12, weight="bold")
        self.mono_font = font.Font(family="Consolas", size=10)
        self.large_font = font.Font(family="Segoe UI", size=13)
    
    def create_widgets(self):
        """Create the main GUI with scrollable interface"""
        
        # Header (Fixed)
        header_frame = tk.Frame(self.root, bg=COLORS['dark'], height=80)
        header_frame.pack(fill=tk.X)
        header_frame.pack_propagate(False)
        
        title_label = tk.Label(
            header_frame,
            text="🎓 Enhanced University Information Assistant",
            font=self.title_font,
            fg='white', bg=COLORS['dark']
        )
        title_label.pack(expand=True)
        
        self.status_indicator = tk.Label(
            header_frame,
            text="🔴 Loading...",
            font=self.normal_font,
            fg=COLORS['warning'], bg=COLORS['dark']
        )
        self.status_indicator.place(relx=0.95, rely=0.5, anchor='e')
        
        # Main scrollable content area
        self.main_scroll = ScrollableFrame(self.root)
        self.main_scroll.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Content container (inside scrollable frame)
        content = self.main_scroll.scrollable_frame
        
        # Query section
        self.create_query_section(content)
        
        # Results section
        self.create_results_section(content)
        
        # Status bar (Fixed)
        status_frame = tk.Frame(self.root, bg=COLORS['dark'], height=35)
        status_frame.pack(fill=tk.X, side=tk.BOTTOM)
        status_frame.pack_propagate(False)
        
        self.status_bar = tk.Label(
            status_frame, textvariable=self.status_var,
            font=self.mono_font, fg='white', bg=COLORS['dark']
        )
        self.status_bar.pack(fill=tk.X, padx=10, pady=5)
    
    def create_query_section(self, parent):
        """Create the query input section"""
        query_frame = tk.LabelFrame(
            parent,
            text="Ask Your Question",
            font=self.subtitle_font,
            bg=COLORS['white'],
            fg=COLORS['text_dark'],
            padx=20,
            pady=15
        )
        query_frame.pack(fill=tk.X, pady=(0, 20))
        
        # Instructions
        instr_label = tk.Label(
            query_frame,
            text="💡 Ask questions about university programs, admissions, policies, facilities, etc.",
            font=self.normal_font,
            fg=COLORS['text_dark'],
            bg=COLORS['white']
        )
        instr_label.pack(anchor=tk.W, pady=(0, 15))
        
        # Query input with better sizing
        self.query_text = scrolledtext.ScrolledText(
            query_frame,
            height=5,
            font=self.normal_font,
            wrap=tk.WORD,
            bg=COLORS['light'],
            relief=tk.SOLID,
            borderwidth=1
        )
        self.query_text.pack(fill=tk.X, pady=(0, 15))
        
        # Buttons frame
        btn_frame = tk.Frame(query_frame, bg=COLORS['white'])
        btn_frame.pack(fill=tk.X, pady=(0, 10))
        
        self.query_btn = tk.Button(
            btn_frame,
            text="🔍 Get Answer",
            command=self.process_query,
            font=self.bold_font,
            bg=COLORS['primary'],
            fg='white',
            relief=tk.FLAT,
            padx=25,
            pady=12,
            state=tk.DISABLED,
            cursor='hand2'
        )
        self.query_btn.pack(side=tk.LEFT, padx=(0, 15))
        
        clear_btn = tk.Button(
            btn_frame,
            text="🧹 Clear",
            command=self.clear_query,
            font=self.bold_font,
            bg=COLORS['secondary'],
            fg='white',
            relief=tk.FLAT,
            padx=25,
            pady=12,
            cursor='hand2'
        )
        clear_btn.pack(side=tk.LEFT)
        
        # Sample queries with better layout
        sample_frame = tk.LabelFrame(
            parent,
            text="Sample Questions",
            font=self.subtitle_font,
            bg=COLORS['white'],
            fg=COLORS['text_dark'],
            padx=20,
            pady=15
        )
        sample_frame.pack(fill=tk.X, pady=(0, 30))
        
        # Create a grid for sample buttons
        samples = [
            "What programs does the university offer?",
            "What are the admission requirements?",
            "What is the fee structure?",
            "What scholarships are available?",
            "What are the campus facilities?"
        ]
        
        for i, sample in enumerate(samples):
            row = i // 2
            col = i % 2
            
            btn = tk.Button(
                sample_frame,
                text=sample,
                command=lambda s=sample: self.set_sample_query(s),
                font=self.normal_font,
                bg=COLORS['info'],
                fg='white',
                relief=tk.FLAT,
                padx=15,
                pady=8,
                cursor='hand2',
                wraplength=300
            )
            btn.grid(row=row, column=col, padx=10, pady=5, sticky='ew')
        
        # Configure grid weights
        sample_frame.grid_columnconfigure(0, weight=1)
        sample_frame.grid_columnconfigure(1, weight=1)
    
    def create_results_section(self, parent):
        """Create the results section with proper scrolling"""
        results_container = tk.Frame(parent, bg=COLORS['light'])
        results_container.pack(fill=tk.BOTH, expand=True)
        
        # Answer section
        answer_frame = tk.LabelFrame(
            results_container,
            text="📝 Answer",
            font=self.subtitle_font,
            bg=COLORS['white'],
            fg=COLORS['text_dark'],
            padx=20,
            pady=15
        )
        answer_frame.pack(fill=tk.X, pady=(0, 20))
        
        # Answer text with scrollbar
        answer_container = tk.Frame(answer_frame, bg=COLORS['white'])
        answer_container.pack(fill=tk.X, pady=(0, 15))
        
        self.answer_text = scrolledtext.ScrolledText(
            answer_container,
            height=12,
            font=self.large_font,
            wrap=tk.WORD,
            bg=COLORS['light'],
            relief=tk.SOLID,
            borderwidth=1,
            state=tk.DISABLED
        )
        self.answer_text.pack(fill=tk.X)
        
        # Action buttons with better spacing
        action_frame = tk.Frame(answer_frame, bg=COLORS['white'])
        action_frame.pack(fill=tk.X, pady=(0, 10))
        
        self.detail_btn = tk.Button(
            action_frame,
            text="📖 Show Detailed Explanation",
            command=self.toggle_detailed_explanation,
            font=self.bold_font,
            bg=COLORS['warning'],
            fg='white',
            relief=tk.FLAT,
            padx=20,
            pady=10,
            cursor='hand2',
            state=tk.DISABLED
        )
        self.detail_btn.pack(side=tk.LEFT, padx=(0, 15))
        
        self.export_btn = tk.Button(
            action_frame,
            text="💾 Export Answer",
            command=self.export_answer,
            font=self.bold_font,
            bg=COLORS['success'],
            fg='white',
            relief=tk.FLAT,
            padx=20,
            pady=10,
            cursor='hand2',
            state=tk.DISABLED
        )
        self.export_btn.pack(side=tk.LEFT)
        
        # Detailed explanation section (initially hidden)
        self.detail_frame = tk.LabelFrame(
            results_container,
            text="📚 Detailed Explanation",
            font=self.subtitle_font,
            bg=COLORS['white'],
            fg=COLORS['text_dark'],
            padx=20,
            pady=15
        )
        
        self.detail_text = scrolledtext.ScrolledText(
            self.detail_frame,
            height=20,
            font=self.large_font,
            wrap=tk.WORD,
            bg=COLORS['bg_secondary'],
            relief=tk.SOLID,
            borderwidth=1,
            state=tk.DISABLED
        )
        self.detail_text.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        
        # Documents section with improved layout
        docs_frame = tk.LabelFrame(
            results_container,
            text="📄 Referenced Documents (Double-Click to Open PDF)",
            font=self.subtitle_font,
            bg=COLORS['white'],
            fg=COLORS['text_dark'],
            padx=20,
            pady=15
        )
        docs_frame.pack(fill=tk.X, pady=(20, 0))
        
        # Create treeview with scrollbars
        tree_container = tk.Frame(docs_frame, bg=COLORS['white'])
        tree_container.pack(fill=tk.X, pady=(0, 10))
        
        # Configure treeview style
        style = ttk.Style()
        style.configure("Custom.Treeview", font=self.normal_font, rowheight=30)
        style.configure("Custom.Treeview.Heading", font=self.bold_font)
        
        self.docs_tree = ttk.Treeview(
            tree_container,
            columns=('filename', 'relevance'),
            show='headings',
            height=8,
            style="Custom.Treeview"
        )
        
        self.docs_tree.heading('filename', text='📄 Document Name')
        self.docs_tree.heading('relevance', text='🎯 Relevance Score')
        
        self.docs_tree.column('filename', width=500, minwidth=300)
        self.docs_tree.column('relevance', width=150, minwidth=100)
        
        # Add scrollbars to treeview
        tree_scroll_y = ttk.Scrollbar(tree_container, orient="vertical", command=self.docs_tree.yview)
        tree_scroll_x = ttk.Scrollbar(tree_container, orient="horizontal", command=self.docs_tree.xview)
        
        self.docs_tree.configure(yscrollcommand=tree_scroll_y.set, xscrollcommand=tree_scroll_x.set)
        
        # Pack treeview and scrollbars
        self.docs_tree.grid(row=0, column=0, sticky='nsew')
        tree_scroll_y.grid(row=0, column=1, sticky='ns')
        tree_scroll_x.grid(row=1, column=0, sticky='ew')
        
        tree_container.grid_rowconfigure(0, weight=1)
        tree_container.grid_columnconfigure(0, weight=1)
        
        # Bind events
        self.docs_tree.bind('<Double-1>', self.open_document)
        self.docs_tree.bind('<Return>', self.open_document)
        
        # Instructions for documents
        docs_instr = tk.Label(
            docs_frame,
            text="💡 Double-click any document to open it with your default PDF/Word viewer",
            font=self.normal_font,
            fg=COLORS['text_light'],
            bg=COLORS['white']
        )
        docs_instr.pack(anchor=tk.W)
        
        # Configure text tags
        self.setup_text_tags()
    
    def setup_text_tags(self):
        """Setup text tags for enhanced formatting"""
        # Answer text tags
        self.answer_text.tag_configure("bold", font=self.bold_font, foreground=COLORS['primary'])
        self.answer_text.tag_configure("heading", font=self.subtitle_font, foreground=COLORS['primary'])
        self.answer_text.tag_configure("normal", font=self.large_font)
        self.answer_text.tag_configure("highlight", background=COLORS['bg_primary'])
        
        # Detail text tags
        self.detail_text.tag_configure("bold", font=self.bold_font, foreground=COLORS['primary'])
        self.detail_text.tag_configure("heading", font=self.subtitle_font, foreground=COLORS['primary'])
        self.detail_text.tag_configure("normal", font=self.large_font)
        self.detail_text.tag_configure("highlight", background=COLORS['bg_primary'])
    
    def initialize_system(self):
        """Initialize system with thread-safe GUI updates"""
        def init_thread():
            try:
                self.root.after(0, lambda: self.update_status("🚀 Loading university documents..."))
                
                documents = self.document_processor.load_all_university_documents()
                
                if not documents:
                    self.root.after(0, lambda: self.update_status("⚠️ No documents found"))
                    self.root.after(0, lambda: messagebox.showwarning(
                        "No Documents",
                        f"No documents found in:\n{UNIVERSITY_DOCS_DIR}\n\nPlease add PDF, DOCX, or text files."
                    ))
                    return
                
                self.root.after(0, lambda: self.update_status(f"📊 Processing {len(documents)} document chunks..."))
                
                self.query_processor.initialize_system(documents)
                self.system_ready = True
                
                self.root.after(0, lambda: self.update_status("✅ System ready! You can now ask questions."))
                self.root.after(0, lambda: self.query_btn.configure(state=tk.NORMAL))
                self.root.after(0, lambda: self.status_indicator.configure(text="🟢 Ready", fg=COLORS['success']))
                
            except Exception as e:
                error_msg = f"Initialization failed: {str(e)}"
                self.root.after(0, lambda: self.update_status(f"❌ {error_msg}"))
                self.root.after(0, lambda: messagebox.showerror("Error", error_msg))
        
        threading.Thread(target=init_thread, daemon=True).start()
    
    def set_sample_query(self, query):
        self.query_text.delete(1.0, tk.END)
        self.query_text.insert(1.0, query)
    
    def clear_query(self):
        self.query_text.delete(1.0, tk.END)
        self.clear_results()
    
    def clear_results(self):
        # Clear main answer
        self.answer_text.config(state=tk.NORMAL)
        self.answer_text.delete(1.0, tk.END)
        self.answer_text.config(state=tk.DISABLED)
        
        # Clear detailed explanation
        self.detail_text.config(state=tk.NORMAL)
        self.detail_text.delete(1.0, tk.END)
        self.detail_text.config(state=tk.DISABLED)
        
        # Hide detailed explanation if visible
        if self.detailed_visible:
            self.detail_frame.pack_forget()
            self.detailed_visible = False
            self.detail_btn.configure(text="📖 Show Detailed Explanation")
        
        # Disable action buttons
        self.detail_btn.configure(state=tk.DISABLED)
        self.export_btn.configure(state=tk.DISABLED)
        
        # Clear documents tree
        for item in self.docs_tree.get_children():
            self.docs_tree.delete(item)
        self.document_paths.clear()
        
        # Reset variables
        self.current_result = None
        self.current_detailed_answer = ""
    
    def process_query(self):
        query = self.query_text.get(1.0, tk.END).strip()
        
        if not query:
            messagebox.showwarning("No Query", "Please enter a question.")
            return
        
        if not self.system_ready:
            messagebox.showwarning("System Not Ready", "Please wait for system to initialize.")
            return
        
        # Disable button during processing
        self.query_btn.configure(state=tk.DISABLED, text="🔄 Processing...")
        self.clear_results()
        
        def query_thread():
            try:
                self.root.after(0, lambda: self.update_status("🤔 Searching documents..."))
                
                result = self.query_processor.process_query(query)
                self.current_result = result
                self.current_detailed_answer = result.detailed_answer
                
                self.root.after(0, lambda: self.display_results(result))
                self.root.after(0, lambda: self.update_status("✅ Query complete"))
                self.root.after(0, lambda: self.query_btn.configure(state=tk.NORMAL, text="🔍 Get Answer"))
                
                # Enable action buttons
                self.root.after(0, lambda: self.detail_btn.configure(state=tk.NORMAL))
                self.root.after(0, lambda: self.export_btn.configure(state=tk.NORMAL))
                
            except Exception as e:
                error_msg = f"Error: {str(e)}"
                self.root.after(0, lambda: self.update_status(f"❌ {error_msg}"))
                self.root.after(0, lambda: messagebox.showerror("Error", error_msg))
                self.root.after(0, lambda: self.query_btn.configure(state=tk.NORMAL, text="🔍 Get Answer"))
        
        threading.Thread(target=query_thread, daemon=True).start()
    
    def display_results(self, result: AnalysisResponse):
        """Display formatted results with enhanced text styling"""
        try:
            # Display main answer
            self.answer_text.config(state=tk.NORMAL)
            self.answer_text.delete(1.0, tk.END)
            
            # Add confidence score
            confidence_text = f"🎯 Confidence Score: {result.confidence_score:.1%}\n\n"
            self.answer_text.insert(tk.END, confidence_text, "heading")
            
            # Format and display answer
            self.insert_formatted_text(self.answer_text, result.answer)
            
            self.answer_text.config(state=tk.DISABLED)
            
            # Populate documents tree
            self.populate_documents_tree(result.sources)
            
        except Exception as e:
            logger.error(f"Error displaying results: {e}")
            self.answer_text.config(state=tk.NORMAL)
            self.answer_text.insert(1.0, f"❌ Error displaying results: {str(e)}")
            self.answer_text.config(state=tk.DISABLED)
    
    def insert_formatted_text(self, text_widget, text):
        """Insert text with formatting"""
        pos = 0
        while pos < len(text):
            start = text.find('**', pos)
            if start == -1:
                text_widget.insert(tk.END, text[pos:], "normal")
                break
            
            if start > pos:
                text_widget.insert(tk.END, text[pos:start], "normal")
            
            end = text.find('**', start + 2)
            if end == -1:
                text_widget.insert(tk.END, text[start:], "normal")
                break
            
            bold_text = text[start + 2:end]
            text_widget.insert(tk.END, bold_text, "bold")
            pos = end + 2
    
    def populate_documents_tree(self, sources):
        """Populate the documents tree with clickable references"""
        print(f"📄 Populating {len(sources)} documents...")
        
        for source in sources:
            filename = source.get('filename', 'Unknown')
            relevance = f"{source.get('relevance', 0):.1%}"
            filepath = source.get('filepath', '')
            
            if filepath and os.path.exists(filepath):
                item = self.docs_tree.insert('', 'end', values=(f"📄 {filename}", relevance))
                self.document_paths[item] = filepath
                print(f"✅ Added: {filename}")
            else:
                item = self.docs_tree.insert('', 'end', values=(f"❌ {filename} (Missing)", relevance))
                print(f"❌ File not found: {filename}")
    
    def toggle_detailed_explanation(self):
        """Toggle the detailed explanation panel"""
        if not self.current_result:
            messagebox.showinfo("No Results", "Please ask a question first.")
            return
        
        if self.detailed_visible:
            # Hide detailed explanation
            self.detail_frame.pack_forget()
            self.detailed_visible = False
            self.detail_btn.configure(text="📖 Show Detailed Explanation", bg=COLORS['warning'])
        else:
            # Show detailed explanation
            self.detail_frame.pack(fill=tk.X, pady=(20, 0))
            self.detailed_visible = True
            self.detail_btn.configure(text="📖 Hide Detailed Explanation", bg=COLORS['danger'])
            
            # Display detailed explanation
            self.display_detailed_explanation()
    
    def display_detailed_explanation(self):
        """Display the detailed explanation"""
        self.detail_text.config(state=tk.NORMAL)
        self.detail_text.delete(1.0, tk.END)
        
        # Add heading
        self.detail_text.insert(tk.END, "📚 COMPREHENSIVE ANALYSIS\n\n", "heading")
        
        # Insert detailed answer
        if self.current_detailed_answer:
            self.insert_formatted_text(self.detail_text, self.current_detailed_answer)
        else:
            self.detail_text.insert(tk.END, "No detailed explanation available.", "normal")
        
        self.detail_text.config(state=tk.DISABLED)
        self.detail_text.see(1.0)  # Scroll to top
    
    def open_document(self, event):
        """Open the selected document"""
        selection = self.docs_tree.selection()
        if not selection:
            messagebox.showinfo("No Selection", "Please select a document to open.")
            return
        
        item = selection[0]
        filepath = self.document_paths.get(item)
        
        if not filepath:
            messagebox.showwarning("No Path", "No file path associated with this document.")
            return
            
        if not os.path.exists(filepath):
            messagebox.showwarning("File Not Found", f"Document not found:\n{filepath}")
            return
        
        try:
            # Cross-platform file opening
            if platform.system() == 'Windows':
                os.startfile(filepath)
            elif platform.system() == 'Darwin':  # macOS
                subprocess.call(['open', filepath])
            else:  # Linux
                subprocess.call(['xdg-open', filepath])
            
            filename = os.path.basename(filepath)
            self.update_status(f"📄 Opened document: {filename}")
            messagebox.showinfo("Document Opened", f"Successfully opened:\n{filename}")
                
        except Exception as e:
            error_msg = f"Could not open document:\n{filepath}\n\nError: {str(e)}"
            messagebox.showerror("Error Opening Document", error_msg)
    
    def export_answer(self):
        """Export the current answer to a text file"""
        if not self.current_result:
            messagebox.showinfo("No Answer", "Please ask a question first.")
            return
        
        try:
            from tkinter import filedialog
            
            filename = filedialog.asksaveasfilename(
                defaultextension=".txt",
                filetypes=[("Text files", "*.txt"), ("All files", "*.*")],
                title="Save Answer As"
            )
            
            if filename:
                with open(filename, 'w', encoding='utf-8') as f:
                    f.write("University Information Assistant - Exported Answer\n")
                    f.write(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                    f.write("="*60 + "\n\n")
                    f.write(f"QUESTION:\n{self.query_text.get(1.0, tk.END).strip()}\n\n")
                    f.write(f"ANSWER:\n{self.current_result.answer}\n\n")
                    
                    if self.current_detailed_answer:
                        f.write(f"DETAILED EXPLANATION:\n{self.current_detailed_answer}\n\n")
                    
                    f.write(f"CONFIDENCE SCORE: {self.current_result.confidence_score:.1%}\n\n")
                    f.write("REFERENCE DOCUMENTS:\n")
                    for ref in self.current_result.document_references:
                        f.write(f"- {ref}\n")
                
                messagebox.showinfo("Export Successful", f"Answer exported to:\n{filename}")
                
        except Exception as e:
            messagebox.showerror("Export Error", f"Error exporting answer:\n{str(e)}")
    
    def update_status(self, message):
        """Update status bar with timestamp"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        self.status_var.set(f"[{timestamp}] {message}")

def main():
    """Main application entry point"""
    root = tk.Tk()
    root.minsize(1400, 900)
    
    app = EnhancedUniversityGUI(root)
    
    def on_closing():
        if messagebox.askokcancel("Quit", "Exit Enhanced University Assistant?"):
            root.destroy()
    
    root.protocol("WM_DELETE_WINDOW", on_closing)
    
    # Center window
    root.update_idletasks()
    x = (root.winfo_screenwidth() // 2) - (root.winfo_width() // 2)
    y = (root.winfo_screenheight() // 2) - (root.winfo_height() // 2)
    root.geometry(f"+{x}+{y}")
    
    print("🎓 Enhanced University Information Assistant Started")
    print(f"📁 Documents directory: {UNIVERSITY_DOCS_DIR}")
    
    root.mainloop()

if __name__ == "__main__":
    main()
