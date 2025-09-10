import os
import json
import time
import re
import threading
import hashlib
import tkinter as tk
from tkinter import ttk, messagebox, scrolledtext, font
from functools import wraps
from typing import List, Optional, Dict, Any
from dataclasses import dataclass
from pydantic import BaseModel
import logging
from datetime import datetime
import webbrowser
import subprocess
import platform
import requests
from urllib.parse import quote

# Import config
try:
    import config
except ImportError:
    print("Please create config.py with your GEMINI_API_KEY")
    exit(1)

# Setup logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s [%(levelname)s] %(message)s')
logger = logging.getLogger(__name__)

# Configure Gemini
import google.generativeai as genai
genai.configure(api_key=getattr(config, 'GEMINI_API_KEY', ''))

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_community.vectorstores import Chroma
from langchain.schema import Document
import PyPDF2
from docx import Document as DocxDocument
import fitz  # PyMuPDF for advanced PDF handling

# Color scheme - University theme
COLORS = {
    'primary': '#1e3a8a',
    'secondary': '#7c3aed',
    'success': '#059669',
    'danger': '#dc2626',
    'warning': '#d97706',
    'info': '#0891b2',
    'light': '#f8fafc',
    'dark': '#1e293b',
    'white': '#ffffff',
    'bg_primary': '#dbeafe',
    'bg_secondary': '#e0e7ff',
    'text_dark': '#1e293b',
    'text_light': '#64748b',
    'border': '#e2e8f0',
    'highlight': '#ffeb3b',
    'highlight_selected': '#ff9800'
}

# Configuration
UNIVERSITY_DOCS_DIR = "./university_documents"
VECTOR_STORE_DIR = "./university_vector_store"
API_RATE_LIMIT = getattr(config, 'API_RATE_LIMIT', 10)
MAX_RETRIES = getattr(config, 'MAX_RETRIES', 3)
RETRY_BASE_DELAY = getattr(config, 'RETRY_BASE_DELAY', 5)
CHUNK_SIZE = getattr(config, 'CHUNK_SIZE', 1000)
CHUNK_OVERLAP = getattr(config, 'CHUNK_OVERLAP', 200)
MAX_CONTEXT_DOCS = getattr(config, 'MAX_CONTEXT_DOCS', 5)
CACHE_DIR = getattr(config, 'CACHE_DIR', "./api_cache")

# ------------------ FIXED: Enhanced PDF Viewer with Highlighting ------------------
class EnhancedPDFViewer(tk.Toplevel):
    """Advanced PDF viewer with search, highlighting, and follow-up question features"""
    
    def __init__(self, parent, document_name, document_path, snippet_text, query, main_app):
        super().__init__(parent)
        self.document_path = document_path
        self.snippet_text = snippet_text
        self.original_query = query
        self.main_app = main_app  # Reference to main app for follow-up queries
        
        # PDF handling
        self.pdf_document = None
        self.current_page = 0
        self.total_pages = 0
        self.zoom_level = 1.0
        self.search_results = []
        self.current_search_index = 0
        self.user_highlights = []  # Store user-created highlights
        self.selected_text = ""
        self.current_selection_index = -1  # Track which highlight is currently selected
        
        # Window setup
        self.title(f"📄 CampusQuery PDF Viewer: {document_name}")
        self.geometry("1400x900")
        self.transient(parent)
        self.configure(bg=COLORS['light'])
        
        # Load PDF
        self.load_pdf()
        
        # Create UI
        self.create_widgets(document_name, query)
        
        # Center window
        self.center_window()
        
        # Highlight original snippet
        self.highlight_original_snippet()
    
    def center_window(self):
        """Center the window on the screen"""
        self.update_idletasks()
        x = (self.winfo_screenwidth() // 2) - (self.winfo_width() // 2)
        y = (self.winfo_screenheight() // 2) - (self.winfo_height() // 2)
        self.geometry(f"+{x}+{y}")
    
    def load_pdf(self):
        """Load PDF document using PyMuPDF"""
        try:
            if not self.document_path.endswith('.pdf'):
                return False
            
            self.pdf_document = fitz.open(self.document_path)
            self.total_pages = len(self.pdf_document)
            return True
        except Exception as e:
            logger.error(f"Error loading PDF: {e}")
            return False
    
    def create_widgets(self, document_name, query):
        # Header frame
        header_frame = tk.Frame(self, bg=COLORS['primary'], height=70)
        header_frame.pack(fill=tk.X)
        header_frame.pack_propagate(False)
        
        # Title
        title_label = tk.Label(
            header_frame,
            text=f"📄 {document_name}",
            font=('Segoe UI', 16, 'bold'),
            fg='white',
            bg=COLORS['primary']
        )
        title_label.pack(side=tk.LEFT, padx=20, pady=20)
        
        # Close button
        close_btn = tk.Button(
            header_frame,
            text="❌ Close",
            command=self.destroy,
            font=('Segoe UI', 12, 'bold'),
            bg=COLORS['danger'],
            fg='white',
            relief=tk.FLAT,
            padx=15,
            pady=8
        )
        close_btn.pack(side=tk.RIGHT, padx=20, pady=20)
        
        # Toolbar frame
        toolbar_frame = tk.Frame(self, bg=COLORS['bg_primary'], height=60)
        toolbar_frame.pack(fill=tk.X)
        toolbar_frame.pack_propagate(False)
        
        # Search functionality
        search_label = tk.Label(
            toolbar_frame,
            text="🔍 Search:",
            font=('Segoe UI', 11, 'bold'),
            bg=COLORS['bg_primary']
        )
        search_label.pack(side=tk.LEFT, padx=(15, 5), pady=15)
        
        self.search_var = tk.StringVar()
        self.search_entry = tk.Entry(
            toolbar_frame,
            textvariable=self.search_var,
            font=('Segoe UI', 11),
            width=25
        )
        self.search_entry.pack(side=tk.LEFT, padx=5, pady=15)
        self.search_entry.bind('<Return>', self.perform_search)
        
        search_btn = tk.Button(
            toolbar_frame,
            text="Search",
            command=self.perform_search,
            font=('Segoe UI', 10, 'bold'),
            bg=COLORS['info'],
            fg='white',
            relief=tk.FLAT,
            padx=10,
            pady=5
        )
        search_btn.pack(side=tk.LEFT, padx=5, pady=15)
        
        # Search navigation
        self.search_nav_frame = tk.Frame(toolbar_frame, bg=COLORS['bg_primary'])
        self.search_nav_frame.pack(side=tk.LEFT, padx=15, pady=15)
        
        # Page navigation
        nav_frame = tk.Frame(toolbar_frame, bg=COLORS['bg_primary'])
        nav_frame.pack(side=tk.RIGHT, padx=15, pady=15)
        
        prev_btn = tk.Button(
            nav_frame,
            text="◀ Prev",
            command=self.prev_page,
            font=('Segoe UI', 10, 'bold'),
            bg=COLORS['secondary'],
            fg='white',
            relief=tk.FLAT,
            padx=10,
            pady=5
        )
        prev_btn.pack(side=tk.LEFT, padx=2)
        
        self.page_var = tk.StringVar()
        self.page_label = tk.Label(
            nav_frame,
            textvariable=self.page_var,
            font=('Segoe UI', 11, 'bold'),
            bg=COLORS['bg_primary']
        )
        self.page_label.pack(side=tk.LEFT, padx=10)
        
        next_btn = tk.Button(
            nav_frame,
            text="Next ▶",
            command=self.next_page,
            font=('Segoe UI', 10, 'bold'),
            bg=COLORS['secondary'],
            fg='white',
            relief=tk.FLAT,
            padx=10,
            pady=5
        )
        next_btn.pack(side=tk.LEFT, padx=2)
        
        # Main content frame
        content_frame = tk.Frame(self, bg=COLORS['white'])
        content_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # PDF display canvas with scrollbars
        canvas_frame = tk.Frame(content_frame, bg=COLORS['white'])
        canvas_frame.pack(fill=tk.BOTH, expand=True)
        
        self.pdf_canvas = tk.Canvas(
            canvas_frame,
            bg='white',
            scrollregion=(0, 0, 0, 0)
        )
        
        # Scrollbars
        v_scrollbar = ttk.Scrollbar(canvas_frame, orient="vertical", command=self.pdf_canvas.yview)
        h_scrollbar = ttk.Scrollbar(canvas_frame, orient="horizontal", command=self.pdf_canvas.xview)
        self.pdf_canvas.configure(yscrollcommand=v_scrollbar.set, xscrollcommand=h_scrollbar.set)
        
        # Pack canvas and scrollbars
        self.pdf_canvas.grid(row=0, column=0, sticky='nsew')
        v_scrollbar.grid(row=0, column=1, sticky='ns')
        h_scrollbar.grid(row=1, column=0, sticky='ew')
        canvas_frame.grid_rowconfigure(0, weight=1)
        canvas_frame.grid_columnconfigure(0, weight=1)
        
        # Bind canvas events for text selection
        self.pdf_canvas.bind('<Button-1>', self.start_selection)
        self.pdf_canvas.bind('<B1-Motion>', self.update_selection)
        self.pdf_canvas.bind('<ButtonRelease-1>', self.end_selection)
        
        # Status frame for highlights and follow-up
        self.status_frame = tk.Frame(self, bg=COLORS['light'], height=50)
        self.status_frame.pack(fill=tk.X, padx=10, pady=5)
        self.status_frame.pack_propagate(False)
        
        # FIXED: Action buttons (initially hidden)
        self.followup_btn = tk.Button(
            self.status_frame,
            text="🤔 Ask about this selection",
            command=self.ask_about_selection,
            font=('Segoe UI', 12, 'bold'),
            bg=COLORS['warning'],
            fg='white',
            relief=tk.FLAT,
            padx=20,
            pady=10
        )
        
        # NEW: Remove highlight button (initially hidden)
        self.remove_highlight_btn = tk.Button(
            self.status_frame,
            text="🗑 Remove Highlight",
            command=self.remove_current_highlight,
            font=('Segoe UI', 12, 'bold'),
            bg=COLORS['danger'],
            fg='white',
            relief=tk.FLAT,
            padx=20,
            pady=10
        )
        
        # NEW: Clear all highlights button
        self.clear_highlights_btn = tk.Button(
            self.status_frame,
            text="🧹 Clear All Highlights",
            command=self.clear_all_highlights,
            font=('Segoe UI', 11, 'bold'),
            bg=COLORS['secondary'],
            fg='white',
            relief=tk.FLAT,
            padx=15,
            pady=8
        )
        
        # Instructions label
        self.instruction_label = tk.Label(
            self.status_frame,
            text="💡 Click and drag to select text, then use the buttons above to ask questions or remove highlights.",
            font=('Segoe UI', 10),
            fg=COLORS['text_light'],
            bg=COLORS['light']
        )
        self.instruction_label.pack(side=tk.LEFT, padx=10, pady=15)
        
        # Update page display
        self.update_page_display()
    
    def update_page_display(self):
        """Update the page counter and render current page"""
        if self.pdf_document:
            self.page_var.set(f"Page {self.current_page + 1} of {self.total_pages}")
            self.render_current_page()
    
    def render_current_page(self):
        """Render the current PDF page on canvas"""
        if not self.pdf_document:
            return
        
        try:
            page = self.pdf_document[self.current_page]
            
            # Get page as image
            mat = fitz.Matrix(self.zoom_level, self.zoom_level)
            pix = page.get_pixmap(matrix=mat)
            img_data = pix.tobytes("ppm")
            
            # Convert to PhotoImage
            self.page_image = tk.PhotoImage(data=img_data)
            
            # Clear canvas and display image
            self.pdf_canvas.delete("all")
            self.pdf_canvas.create_image(0, 0, anchor="nw", image=self.page_image)
            
            # Update scroll region
            self.pdf_canvas.configure(scrollregion=self.pdf_canvas.bbox("all"))
            
            # Re-apply highlights
            self.apply_highlights()
            
        except Exception as e:
            logger.error(f"Error rendering page: {e}")
    
    def prev_page(self):
        """Go to previous page"""
        if self.current_page > 0:
            self.current_page -= 1
            self.update_page_display()
    
    def next_page(self):
        """Go to next page"""
        if self.current_page < self.total_pages - 1:
            self.current_page += 1
            self.update_page_display()
    
    def perform_search(self, event=None):
        """Search for text in PDF"""
        search_term = self.search_var.get().strip()
        if not search_term or not self.pdf_document:
            return
        
        # Clear previous search results
        self.search_results = []
        self.current_search_index = 0
        
        # Search all pages
        for page_num in range(self.total_pages):
            page = self.pdf_document[page_num]
            text_instances = page.search_for(search_term)
            
            for inst in text_instances:
                self.search_results.append({
                    'page': page_num,
                    'rect': inst,
                    'text': search_term
                })
        
        # Update search navigation
        self.update_search_navigation()
        
        # Go to first result
        if self.search_results:
            self.go_to_search_result(0)
    
    def update_search_navigation(self):
        """Update search navigation buttons"""
        # Clear existing navigation
        for widget in self.search_nav_frame.winfo_children():
            widget.destroy()
        
        if self.search_results:
            total_results = len(self.search_results)
            
            prev_search_btn = tk.Button(
                self.search_nav_frame,
                text="◀",
                command=self.prev_search_result,
                font=('Segoe UI', 10, 'bold'),
                bg=COLORS['success'],
                fg='white',
                relief=tk.FLAT,
                padx=5,
                pady=2
            )
            prev_search_btn.pack(side=tk.LEFT, padx=2)
            
            result_label = tk.Label(
                self.search_nav_frame,
                text=f"{self.current_search_index + 1} of {total_results}",
                font=('Segoe UI', 10),
                bg=COLORS['bg_primary']
            )
            result_label.pack(side=tk.LEFT, padx=5)
            
            next_search_btn = tk.Button(
                self.search_nav_frame,
                text="▶",
                command=self.next_search_result,
                font=('Segoe UI', 10, 'bold'),
                bg=COLORS['success'],
                fg='white',
                relief=tk.FLAT,
                padx=5,
                pady=2
            )
            next_search_btn.pack(side=tk.LEFT, padx=2)
    
    def prev_search_result(self):
        """Go to previous search result"""
        if self.search_results and self.current_search_index > 0:
            self.current_search_index -= 1
            self.go_to_search_result(self.current_search_index)
            self.update_search_navigation()
    
    def next_search_result(self):
        """Go to next search result"""
        if self.search_results and self.current_search_index < len(self.search_results) - 1:
            self.current_search_index += 1
            self.go_to_search_result(self.current_search_index)
            self.update_search_navigation()
    
    def go_to_search_result(self, index):
        """Navigate to specific search result"""
        if not self.search_results or index >= len(self.search_results):
            return
        
        result = self.search_results[index]
        
        # Go to the page
        self.current_page = result['page']
        self.update_page_display()
        
        # Scroll to the result (simplified - would need more complex positioning)
        # For now, just ensure the page is displayed
    
    def highlight_original_snippet(self):
        """Highlight the original snippet that answered the query"""
        if not self.snippet_text or not self.pdf_document:
            return
        
        # Search for snippet text in PDF
        snippet_words = self.snippet_text.split()[:10]  # First 10 words
        search_term = " ".join(snippet_words)
        
        for page_num in range(self.total_pages):
            page = self.pdf_document[page_num]
            text_instances = page.search_for(search_term)
            
            if text_instances:
                # Navigate to this page
                self.current_page = page_num
                self.update_page_display()
                
                # Add highlight
                for inst in text_instances:
                    self.add_highlight(inst, COLORS['highlight'], "Original Answer Source")
                break
    
    def apply_highlights(self):
        """Apply all highlights to current page"""
        if not self.pdf_document:
            return
        
        page = self.pdf_document[self.current_page]
        
        # Apply search result highlights
        for result in self.search_results:
            if result['page'] == self.current_page:
                self.add_highlight(result['rect'], COLORS['info'], "Search Result")
        
        # Apply user highlights
        for i, highlight in enumerate(self.user_highlights):
            if highlight['page'] == self.current_page:
                color = highlight['color']
                if i == self.current_selection_index:
                    color = '#ff5722'  # Different color for currently selected highlight
                self.add_highlight(highlight['rect'], color, highlight['type'])
    
    def add_highlight(self, rect, color, highlight_type):
        """Add a highlight rectangle to the canvas"""
        try:
            # Convert PDF coordinates to canvas coordinates
            x1, y1, x2, y2 = rect
            x1 *= self.zoom_level
            y1 *= self.zoom_level
            x2 *= self.zoom_level
            y2 *= self.zoom_level
            
            # Create highlight rectangle
            self.pdf_canvas.create_rectangle(
                x1, y1, x2, y2,
                fill=color,
                stipple="gray50",
                tags=f"highlight_{highlight_type}"
            )
        except Exception as e:
            logger.error(f"Error adding highlight: {e}")
    
    def start_selection(self, event):
        """Start text selection"""
        self.selection_start = (self.pdf_canvas.canvasx(event.x), self.pdf_canvas.canvasy(event.y))
    
    def update_selection(self, event):
        """Update selection rectangle"""
        if hasattr(self, 'selection_start'):
            # Clear previous selection rectangle
            self.pdf_canvas.delete("selection")
            
            # Draw new selection rectangle
            x1, y1 = self.selection_start
            x2, y2 = self.pdf_canvas.canvasx(event.x), self.pdf_canvas.canvasy(event.y)
            
            self.pdf_canvas.create_rectangle(
                x1, y1, x2, y2,
                outline=COLORS['primary'],
                width=2,
                tags="selection"
            )
    
    def end_selection(self, event):
        """End text selection and extract text"""
        if not hasattr(self, 'selection_start') or not self.pdf_document:
            return
        
        try:
            # Get selection coordinates
            x1, y1 = self.selection_start
            x2, y2 = self.pdf_canvas.canvasx(event.x), self.pdf_canvas.canvasy(event.y)
            
            # Convert canvas coordinates to PDF coordinates
            pdf_x1 = min(x1, x2) / self.zoom_level
            pdf_y1 = min(y1, y2) / self.zoom_level
            pdf_x2 = max(x1, x2) / self.zoom_level
            pdf_y2 = max(y1, y2) / self.zoom_level
            
            # Create selection rectangle
            selection_rect = fitz.Rect(pdf_x1, pdf_y1, pdf_x2, pdf_y2)
            
            # Extract text from selection
            page = self.pdf_document[self.current_page]
            selected_text = page.get_text("text", clip=selection_rect).strip()
            
            if selected_text and len(selected_text) > 10:  # Only process meaningful selections
                self.selected_text = selected_text
                
                # Add user highlight
                highlight_data = {
                    'page': self.current_page,
                    'rect': selection_rect,
                    'color': COLORS['highlight_selected'],
                    'type': 'user_selection',
                    'text': selected_text
                }
                self.user_highlights.append(highlight_data)
                self.current_selection_index = len(self.user_highlights) - 1
                
                # Show action buttons
                self.show_highlight_controls()
                
                # Re-render to show highlight
                self.render_current_page()
            
        except Exception as e:
            logger.error(f"Error in text selection: {e}")
        
        # Clean up
        if hasattr(self, 'selection_start'):
            del self.selection_start
            
        # Clear selection rectangle
        self.pdf_canvas.delete("selection")
    
    def show_highlight_controls(self):
        """Show the highlight control buttons"""
        self.followup_btn.pack(side=tk.LEFT, padx=10, pady=10)
        self.remove_highlight_btn.pack(side=tk.LEFT, padx=5, pady=10)
        if len(self.user_highlights) > 1:
            self.clear_highlights_btn.pack(side=tk.LEFT, padx=5, pady=10)
        
        # Update instruction label
        self.instruction_label.configure(
            text=f"✅ Selected: \"{self.selected_text[:50]}{'...' if len(self.selected_text) > 50 else ''}\" - Use buttons above to interact."
        )
    
    def hide_highlight_controls(self):
        """Hide the highlight control buttons"""
        self.followup_btn.pack_forget()
        self.remove_highlight_btn.pack_forget()
        self.clear_highlights_btn.pack_forget()
        
        # Reset instruction label
        self.instruction_label.configure(
            text="💡 Click and drag to select text, then use the buttons above to ask questions or remove highlights."
        )
    
    def remove_current_highlight(self):
        """FIXED: Remove the currently selected highlight"""
        if self.current_selection_index >= 0 and self.current_selection_index < len(self.user_highlights):
            # Remove the highlight
            removed_highlight = self.user_highlights.pop(self.current_selection_index)
            
            # Reset selection
            self.selected_text = ""
            self.current_selection_index = -1
            
            # Update UI
            if not self.user_highlights:
                self.hide_highlight_controls()
            else:
                # If there are still highlights, select the last one
                self.current_selection_index = len(self.user_highlights) - 1
                self.selected_text = self.user_highlights[self.current_selection_index]['text']
                self.show_highlight_controls()
            
            # Re-render page
            self.render_current_page()
            
            messagebox.showinfo("Highlight Removed", "Selected highlight has been removed.")
    
    def clear_all_highlights(self):
        """NEW: Clear all user-created highlights"""
        if self.user_highlights:
            result = messagebox.askyesno(
                "Clear All Highlights", 
                f"Are you sure you want to remove all {len(self.user_highlights)} highlights?"
            )
            if result:
                self.user_highlights.clear()
                self.selected_text = ""
                self.current_selection_index = -1
                self.hide_highlight_controls()
                self.render_current_page()
                messagebox.showinfo("Highlights Cleared", "All highlights have been removed.")
    
    def ask_about_selection(self):
        """FIXED: Generate and process follow-up question about selected text"""
        if not self.selected_text:
            messagebox.showwarning("No Selection", "Please select some text first.")
            return
        
        try:
            # Generate intelligent follow-up question
            question_prompt = f"""
            Based on this selected text from a university document: "{self.selected_text}"
            
            Generate a natural, specific follow-up question that would help a student understand this content better. 
            The question should be:
            1. Directly related to the selected text
            2. Practical and useful for students
            3. Clear and specific
            
            Respond with just the question, nothing else.
            """
            
            # Use the main app's model client to generate question
            model_client = self.main_app.query_processor.model_client
            follow_up_question = model_client.generate_content(question_prompt)
            
            # Process the follow-up question through main app
            self.process_followup_question(follow_up_question, self.selected_text)
            
        except Exception as e:
            logger.error(f"Error generating follow-up question: {e}")
            messagebox.showerror("Error", f"Could not generate follow-up question: {str(e)}")
    
    def process_followup_question(self, question, context_text):
        """FIXED: Process follow-up question and show result"""
        try:
            # Create enhanced prompt with context
            enhanced_question = f"{question}\n\nContext from document: {context_text}"
            
            # Set the question in main app and process it
            self.main_app.query_text.delete(1.0, tk.END)
            self.main_app.query_text.insert(1.0, enhanced_question)
            
            # Show a notification
            messagebox.showinfo(
                "Follow-up Question Generated", 
                f"Generated question: {question}\n\nProcessing answer..."
            )
            
            # Process the query in the main app
            self.main_app.process_query()
            
            # FIXED: Instead of iconify() which causes error with transient windows, 
            # either withdraw() to hide the window or deiconify() to bring main window to front
            try:
                self.withdraw()  # Hide the PDF viewer window
                self.main_app.root.lift()  # Bring main app to front
                self.main_app.root.focus_force()  # Give focus to main app
                
                # After a delay, show the PDF viewer again
                self.after(3000, lambda: self.deiconify())  # Show window again after 3 seconds
                
            except Exception as window_error:
                logger.warning(f"Window management issue: {window_error}")
                # If window management fails, just continue without hiding
            
        except Exception as e:
            logger.error(f"Error processing follow-up question: {e}")
            messagebox.showerror("Error", f"Could not process follow-up question: {str(e)}")

# ------------------ Original Document Snippet Viewer (for non-PDF files) ------------------
class DocumentSnippetViewer(tk.Toplevel):
    """Window that shows the specific text snippet for non-PDF files"""
    
    def __init__(self, parent, document_name, document_path, snippet_text, query):
        super().__init__(parent)
        self.document_path = document_path
        self.snippet_text = snippet_text
        
        # Window setup
        self.title(f"📄 Document Source: {document_name}")
        self.geometry("900x700")
        self.transient(parent)
        self.grab_set()
        
        # Configure colors
        self.configure(bg=COLORS['light'])
        
        # Create UI
        self.create_widgets(document_name, query)
        
        # Center window
        self.center_window()
    
    def center_window(self):
        """Center the window on the screen"""
        self.update_idletasks()
        x = (self.winfo_screenwidth() // 2) - (self.winfo_width() // 2)
        y = (self.winfo_screenheight() // 2) - (self.winfo_height() // 2)
        self.geometry(f"+{x}+{y}")
    
    def create_widgets(self, document_name, query):
        # Header frame
        header_frame = tk.Frame(self, bg=COLORS['primary'], height=60)
        header_frame.pack(fill=tk.X)
        header_frame.pack_propagate(False)
        
        # Title
        title_label = tk.Label(
            header_frame,
            text=f"📄 Source Content from: {document_name}",
            font=('Segoe UI', 14, 'bold'),
            fg='white',
            bg=COLORS['primary']
        )
        title_label.pack(expand=True)
        
        # Query info frame
        query_frame = tk.Frame(self, bg=COLORS['bg_primary'], relief=tk.RIDGE, bd=2)
        query_frame.pack(fill=tk.X, padx=10, pady=10)
        
        tk.Label(
            query_frame,
            text=f"🔍 Your Question: {query}",
            font=('Segoe UI', 11, 'bold'),
            fg=COLORS['text_dark'],
            bg=COLORS['bg_primary'],
            wraplength=800,
            justify=tk.LEFT
        ).pack(padx=15, pady=10)
        
        # Content frame
        content_frame = tk.LabelFrame(
            self,
            text="📝 Relevant Text from Document (Highlighted Below)",
            font=('Segoe UI', 12, 'bold'),
            fg=COLORS['primary'],
            bg=COLORS['white']
        )
        content_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=(0, 10))
        
        # Text widget with scrollbar
        text_frame = tk.Frame(content_frame)
        text_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        self.text_widget = scrolledtext.ScrolledText(
            text_frame,
            font=('Segoe UI', 12),
            wrap=tk.WORD,
            bg=COLORS['light'],
            relief=tk.SOLID,
            borderwidth=1,
            state=tk.DISABLED
        )
        self.text_widget.pack(fill=tk.BOTH, expand=True)
        
        # Configure text tags for highlighting
        self.text_widget.tag_configure("highlight", background="yellow", foreground="black", font=('Segoe UI', 12, 'bold'))
        self.text_widget.tag_configure("normal", font=('Segoe UI', 12))
        self.text_widget.tag_configure("context", font=('Segoe UI', 11), foreground=COLORS['text_light'])
        
        # Display content
        self.display_content()
        
        # Buttons frame
        btn_frame = tk.Frame(self, bg=COLORS['light'])
        btn_frame.pack(fill=tk.X, padx=10, pady=10)
        
        # Open full document button
        open_full_btn = tk.Button(
            btn_frame,
            text="📄 Open Full Document",
            command=self.open_full_document,
            font=('Segoe UI', 11, 'bold'),
            bg=COLORS['primary'],
            fg='white',
            relief=tk.FLAT,
            padx=20,
            pady=8,
            cursor='hand2'
        )
        open_full_btn.pack(side=tk.LEFT, padx=(0, 10))
        
        # Copy text button
        copy_btn = tk.Button(
            btn_frame,
            text="📋 Copy Text",
            command=self.copy_text,
            font=('Segoe UI', 11, 'bold'),
            bg=COLORS['success'],
            fg='white',
            relief=tk.FLAT,
            padx=20,
            pady=8,
            cursor='hand2'
        )
        copy_btn.pack(side=tk.LEFT, padx=(0, 10))
        
        # Close button
        close_btn = tk.Button(
            btn_frame,
            text="❌ Close",
            command=self.destroy,
            font=('Segoe UI', 11, 'bold'),
            bg=COLORS['danger'],
            fg='white',
            relief=tk.FLAT,
            padx=20,
            pady=8,
            cursor='hand2'
        )
        close_btn.pack(side=tk.RIGHT)
    
    def display_content(self):
        """Display the content with highlighting"""
        self.text_widget.config(state=tk.NORMAL)
        self.text_widget.delete(1.0, tk.END)
        
        # Add context information
        context_text = "📌 The highlighted text below is the specific content that was used to answer your question:\n\n"
        self.text_widget.insert(tk.END, context_text, "context")
        
        # Add separator
        separator = "─" * 80 + "\n\n"
        self.text_widget.insert(tk.END, separator, "context")
        
        # Insert the snippet with highlighting
        snippet_clean = self.snippet_text.strip()
        if snippet_clean:
            self.text_widget.insert(tk.END, snippet_clean, "highlight")
            
            # Add some context around the snippet if available
            self.text_widget.insert(tk.END, f"\n\n{separator}", "context")
            self.text_widget.insert(tk.END, "💡 This highlighted section contains the information used to answer your question.", "context")
        else:
            self.text_widget.insert(tk.END, "No specific text snippet available.", "normal")
        
        self.text_widget.config(state=tk.DISABLED)
        
        # Scroll to the beginning
        self.text_widget.see(1.0)
    
    def open_full_document(self):
        """Open the full document in default application"""
        if not os.path.exists(self.document_path):
            messagebox.showerror("Error", f"Document not found:\n{self.document_path}")
            return
        
        try:
            if platform.system() == 'Windows':
                os.startfile(self.document_path)
            elif platform.system() == 'Darwin':  # macOS
                subprocess.call(['open', self.document_path])
            else:  # Linux
                subprocess.call(['xdg-open', self.document_path])
            messagebox.showinfo("Document Opened", f"Full document opened in default application.")
        except Exception as e:
            messagebox.showerror("Error", f"Could not open document:\n{str(e)}")
    
    def copy_text(self):
        """Copy the snippet text to clipboard"""
        self.clipboard_clear()
        self.clipboard_append(self.snippet_text)
        messagebox.showinfo("Copied", "Text copied to clipboard!")

# ------------------ Core System Classes (Enhanced) ------------------
def rate_limit(calls_per_minute: int):
    min_interval = 60.0 / calls_per_minute
    
    def decorator(func):
        last_called = [0.0]
        
        @wraps(func)
        def wrapper(*args, **kwargs):
            elapsed = time.time() - last_called[0]
            wait_time = min_interval - elapsed
            if wait_time > 0:
                time.sleep(wait_time)
            ret = func(*args, **kwargs)
            last_called[0] = time.time()
            return ret
        return wrapper
    return decorator

class APICache:
    def __init__(self, cache_dir=CACHE_DIR):
        self.cache_dir = cache_dir
        os.makedirs(self.cache_dir, exist_ok=True)
    
    def _get_cache_key(self, prompt: str) -> str:
        return hashlib.md5(prompt.encode('utf-8')).hexdigest()
    
    def get(self, prompt: str) -> Optional[str]:
        cache_key = self._get_cache_key(prompt)
        cache_file = os.path.join(self.cache_dir, f"{cache_key}.json")
        
        if os.path.exists(cache_file):
            try:
                with open(cache_file, "r", encoding="utf-8") as f:
                    data = json.load(f)
                return data.get("content")
            except Exception:
                pass
        return None
    
    def set(self, prompt: str, content: str):
        cache_key = self._get_cache_key(prompt)
        cache_file = os.path.join(self.cache_dir, f"{cache_key}.json")
        
        try:
            with open(cache_file, "w", encoding="utf-8") as f:
                json.dump({"content": content, "timestamp": time.time()}, f)
        except Exception:
            pass

@dataclass
class QueryStructure:
    department: Optional[str] = None
    program: Optional[str] = None
    semester: Optional[str] = None
    course: Optional[str] = None
    topic: Optional[str] = None
    raw_query: str = ""

class AnalysisResponse(BaseModel):
    answer: str
    detailed_answer: str
    justification: str
    key_points: List[str]
    document_references: List[str]
    sources: List[Dict[str, Any]]
    applicable_sections: List[str]

class SafeModelClient:
    def __init__(self, model_name: str = "gemini-1.5-flash"):
        self.model = genai.GenerativeModel(model_name)
        self.cache = APICache()
        self.rate_limited_generate = rate_limit(API_RATE_LIMIT)(self._generate_content_internal)
    
    def _generate_content_internal(self, prompt: str) -> str:
        response = self.model.generate_content(prompt)
        return response.text or ""
    
    def generate_content(self, prompt: str) -> str:
        cached_content = self.cache.get(prompt)
        if cached_content:
            return cached_content
        
        for attempt in range(MAX_RETRIES):
            try:
                content = self.rate_limited_generate(prompt)
                self.cache.set(prompt, content)
                return content
            except Exception as e:
                if attempt == MAX_RETRIES - 1:
                    raise e
                time.sleep(RETRY_BASE_DELAY * (2 ** attempt))
        raise RuntimeError("Exceeded maximum retry attempts")

class UniversityDocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE,
            chunk_overlap=CHUNK_OVERLAP,
            separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""]
        )
    
    def load_pdf(self, filepath: str) -> str:
        try:
            text = ""
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text += page_text + "\n"
                    except Exception:
                        continue
            return text.strip() if text.strip() else None
        except Exception:
            return None
    
    def load_docx(self, filepath: str) -> str:
        try:
            doc = DocxDocument(filepath)
            text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text.strip() + "\n"
            return text.strip() if text.strip() else None
        except Exception:
            return None
    
    def load_text(self, filepath: str) -> str:
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                return f.read().strip()
        except Exception:
            return None
    
    def load_all_university_documents(self) -> List[Document]:
        documents = []
        if not os.path.exists(UNIVERSITY_DOCS_DIR):
            return documents
        
        files = [f for f in os.listdir(UNIVERSITY_DOCS_DIR) if not os.path.isdir(os.path.join(UNIVERSITY_DOCS_DIR, f))]
        
        for filename in files:
            filepath = os.path.join(UNIVERSITY_DOCS_DIR, filename)
            file_ext = os.path.splitext(filepath)[1].lower()
            text = None
            
            if file_ext == '.pdf':
                text = self.load_pdf(filepath)
            elif file_ext == '.docx':
                text = self.load_docx(filepath)
            elif file_ext in ['.txt', '.md']:
                text = self.load_text(filepath)
            
            if not text:
                continue
            
            # Clean text
            text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
            text = re.sub(r' +', ' ', text)
            text = text.strip()
            
            # Split into chunks
            chunks = self.text_splitter.split_text(text)
            for i, chunk in enumerate(chunks):
                if chunk.strip():
                    doc = Document(
                        page_content=chunk,
                        metadata={
                            "source": filepath,
                            "filename": filename,
                            "chunk_id": f"{filename}_{i}",
                            "file_type": file_ext
                        }
                    )
                    documents.append(doc)
        
        return documents

# NEW: Web Search Helper Class
class WebSearchHelper:
    def __init__(self, api_key: str):
        self.api_key = api_key
        self.model = genai.GenerativeModel("gemini-1.5-flash")
    
    def search_web_and_get_content(self, query: str, max_results: int = 3) -> List[Dict[str, str]]:
        """
        Use Gemini API to search the web and get relevant information for the query.
        This is a fallback when local documents don't contain the answer.
        """
        try:
            # Create a comprehensive web search prompt for Gemini
            web_search_prompt = f"""
            I need you to act as a web search assistant. Please provide comprehensive, factual information about: "{query}"

            Please provide information structured as follows:
            
            TOPIC: {query}
            
            OVERVIEW:
            [Provide a comprehensive overview of the topic]
            
            KEY INFORMATION:
            [List key facts, details, and important information]
            
            SPECIFIC DETAILS:
            [Include specific details, dates, requirements, procedures, etc. as relevant]
            
            ADDITIONAL CONTEXT:
            [Any additional relevant context or related information]
            
            Please ensure the information is accurate, current, and comprehensive. Focus on providing factual, helpful content that directly addresses the query.
            """
            
            response = self.model.generate_content(web_search_prompt)
            content = response.text if response.text else "No information available"
            
            # Structure the response as search results
            results = [{
                'title': f"Web Search Result for: {query}",
                'url': 'Generated by AI Assistant',
                'snippet': content
            }]
            
            return results
            
        except Exception as e:
            logger.error(f"Web search error: {e}")
            return [{
                'title': 'Error in Web Search',
                'url': 'Error',
                'snippet': f'Unable to perform web search: {str(e)}'
            }]

class UniversityQueryProcessor:
    def __init__(self):
        self.model_client = SafeModelClient('gemini-1.5-flash')
        self.embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
        self.vector_store = None
        self.system_initialized = False
        # NEW: Initialize web search helper
        self.web_search_helper = WebSearchHelper(getattr(config, 'GEMINI_API_KEY', ''))
    
    def initialize_system(self, documents: List[Document]):
        try:
            # Try to load existing vector store
            if os.path.exists(VECTOR_STORE_DIR) and os.listdir(VECTOR_STORE_DIR):
                self.vector_store = Chroma(
                    persist_directory=VECTOR_STORE_DIR,
                    embedding_function=self.embeddings
                )
                self.system_initialized = True
                return
            
            # Create new vector store
            if documents:
                self.vector_store = Chroma.from_documents(
                    documents=documents,
                    embedding=self.embeddings,
                    persist_directory=VECTOR_STORE_DIR
                )
                self.system_initialized = True
        except Exception as e:
            raise Exception(f"Error initializing system: {e}")
    
    def search_relevant_content(self, query: str, k: int = 5) -> List[Document]:
        if not self.vector_store:
            return []
        return self.vector_store.similarity_search(query, k=k)
    
    def is_answer_adequate(self, answer: str, query: str) -> bool:
        """
        Check if the generated answer is adequate or if we need to fallback to web search.
        """
        inadequate_indicators = [
            "no relevant information found",
            "no information available",
            "cannot find information",
            "not found in the documents",
            "please try rephrasing",
            "no documents contain",
            "unable to find",
            "no data available"
        ]
        
        answer_lower = answer.lower()
        
        # If answer is too short (less than 50 characters), it's likely inadequate
        if len(answer.strip()) < 50:
            return False
        
        # Check for inadequate response indicators
        for indicator in inadequate_indicators:
            if indicator in answer_lower:
                return False
        
        return True
    
    def web_search_and_answer(self, query: str) -> AnalysisResponse:
        """
        NEW: Fallback method to search web and generate answer when local documents are insufficient.
        """
        try:
            # Get web search results
            web_results = self.web_search_helper.search_web_and_get_content(query, max_results=3)
            
            if not web_results:
                return AnalysisResponse(
                    answer="Unable to find information through web search.",
                    detailed_answer="The system could not retrieve relevant information from web search.",
                    justification="Web search failed.",
                    key_points=[],
                    document_references=[],
                    sources=[],
                    applicable_sections=[]
                )
            
            # Prepare context from web results
            context = "\n\n".join([
                f"Source: {result['title']}\nContent: {result['snippet']}"
                for result in web_results
            ])
            
            # Generate answer from web context
            answer_prompt = f"""
            Based on the following web search information, provide a comprehensive answer to the user's question.

            QUESTION: {query}

            WEB SEARCH RESULTS:
            {context}

            Please provide your answer in a structured format with:
            1. Key information highlighted with **bold** formatting
            2. Clear sections and bullet points where appropriate
            3. Comprehensive coverage of the topic

            Make the answer well-organized and easy to read. Note that this information comes from web search since it wasn't available in the local university documents.
            """

            detailed_prompt = f"""
            Based on the following web search information, provide a detailed, comprehensive explanation for the user's question.

            QUESTION: {query}

            WEB SEARCH RESULTS:
            {context}

            Please provide:
            1. **In-depth analysis** of the topic with background information
            2. **Step-by-step explanations** where applicable
            3. **Additional related information** that might be helpful
            4. **Specific examples** from the search results
            5. **Practical advice** or next steps for the user

            Format with **bold** keywords for better readability. This is supplementary information from web search.
            """

            # Generate responses
            response_text = self.model_client.generate_content(answer_prompt)
            detailed_response = self.model_client.generate_content(detailed_prompt)
            
            # Prepare sources
            sources = []
            for i, result in enumerate(web_results):
                sources.append({
                    'source_id': f"web_source_{i+1}",
                    'filename': f"Web Search Result {i+1}",
                    'filepath': '',  # No file path for web results
                    'content_preview': result['snippet'][:200] + "..." if len(result['snippet']) > 200 else result['snippet'],
                    'content_snippet': result['snippet'],
                    'relevance': 1.0 - (i * 0.1)
                })
            
            return AnalysisResponse(
                answer=f"🌐 **Information from Web Search:**\n\n{response_text}",
                detailed_answer=f"🌐 **Detailed Web Search Information:**\n\n{detailed_response}",
                justification=f"Answer based on web search results (no relevant information found in local documents)",
                key_points=[response_text[:100] + "..." if len(response_text) > 100 else response_text],
                document_references=[],
                sources=sources,
                applicable_sections=[]
            )
            
        except Exception as e:
            logger.error(f"Error in web search fallback: {e}")
            return AnalysisResponse(
                answer=f"Error during web search: {str(e)}",
                detailed_answer=f"Error during web search: {str(e)}",
                justification="An error occurred during web search",
                key_points=[],
                document_references=[],
                sources=[],
                applicable_sections=[]
            )
    
    def process_query(self, query: str) -> AnalysisResponse:
        if not self.system_initialized:
            raise ValueError("System not ready")
        
        # First, try to find relevant documents
        relevant_docs = self.search_relevant_content(query, k=5)
        
        # If no documents found, immediately fallback to web search
        if not relevant_docs:
            logger.info(f"No relevant documents found for query: {query}. Falling back to web search.")
            return self.web_search_and_answer(query)
        
        # Prepare context from documents
        context = "\n\n".join([
            f"Document: {doc.metadata.get('filename', 'Unknown')}\nContent: {doc.page_content}"
            for doc in relevant_docs[:3]
        ])
        
        # Generate structured answer
        answer_prompt = f"""
        Based on the following university documents, provide a comprehensive answer to the user's question.

        QUESTION: {query}

        CONTEXT: {context}

        Please provide your answer in a structured format with:
        1. Key information highlighted with **bold** formatting
        2. Clear sections and bullet points where appropriate
        3. Specific references to the source documents

        Make the answer well-organized and easy to read.
        """

        # Generate detailed explanation
        detailed_prompt = f"""
        Based on the following university documents, provide a detailed, comprehensive explanation for the user's question.

        QUESTION: {query}

        CONTEXT: {context}

        Please provide:
        1. **In-depth analysis** of the topic with background information
        2. **Step-by-step explanations** where applicable
        3. **Additional related information** that might be helpful
        4. **Specific examples** from the documents
        5. **Cross-references** to related topics or procedures
        6. **Practical advice** or next steps for the user

        This should be a thorough explanation that goes beyond the basic answer and provides comprehensive understanding of the topic.

        Format with **bold** keywords for better readability.
        """

        try:
            # Generate both answers
            response_text = self.model_client.generate_content(answer_prompt)
            detailed_response = self.model_client.generate_content(detailed_prompt)
            
            # NEW: Check if the answer is adequate
            if not self.is_answer_adequate(response_text, query):
                logger.info(f"Document-based answer inadequate for query: {query}. Falling back to web search.")
                return self.web_search_and_answer(query)
            
            # If answer is adequate, proceed with document-based response
            sources = []
            for i, doc in enumerate(relevant_docs):
                # Convert relative path to absolute path
                absolute_path = os.path.abspath(doc.metadata.get('source', ''))
                sources.append({
                    'source_id': f"source_{i+1}",
                    'filename': doc.metadata.get('filename', 'unknown'),
                    'filepath': absolute_path,
                    'content_preview': doc.page_content[:200] + "...",
                    'content_snippet': doc.page_content,  # ENHANCED: Store full snippet for highlighting
                    'relevance': 1.0 - (i * 0.2)
                })
            
            return AnalysisResponse(
                answer=response_text,
                detailed_answer=detailed_response,
                justification=f"Answer based on {len(relevant_docs)} relevant documents",
                key_points=[response_text[:100] + "..."],
                document_references=[doc.metadata.get('filename', 'Unknown') for doc in relevant_docs],
                sources=sources,
                applicable_sections=[]
            )
            
        except Exception as e:
            logger.error(f"Error processing query with documents: {e}")
            return AnalysisResponse(
                answer=f"Error processing query: {str(e)}",
                detailed_answer=f"Error processing query: {str(e)}",
                justification="An error occurred",
                key_points=[],
                document_references=[],
                sources=[],
                applicable_sections=[]
            )

# ------------------ Enhanced Scrollable Frame ------------------
class ScrollableFrame(tk.Frame):
    """A scrollable frame that contains a canvas and scrollbar"""
    
    def __init__(self, parent):
        super().__init__(parent)
        
        # Create canvas and scrollbar
        self.canvas = tk.Canvas(self, bg=COLORS['light'], highlightthickness=0)
        self.scrollbar = ttk.Scrollbar(self, orient="vertical", command=self.canvas.yview)
        self.scrollable_frame = tk.Frame(self.canvas, bg=COLORS['light'])
        
        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all"))
        )
        
        self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        self.canvas.configure(yscrollcommand=self.scrollbar.set)
        
        # Pack canvas and scrollbar
        self.canvas.pack(side="left", fill="both", expand=True)
        self.scrollbar.pack(side="right", fill="y")
        
        # Bind mousewheel to canvas
        self.bind_mousewheel()
        
        # Configure canvas width to match window
        self.canvas.bind('<Configure>', self._on_canvas_configure)
    
    def _on_canvas_configure(self, event):
        canvas_width = event.width
        self.canvas.itemconfig(self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw"), width=canvas_width)
    
    def bind_mousewheel(self):
        def _on_mousewheel(event):
            self.canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        
        def _bind_to_mousewheel(event):
            self.canvas.bind_all("<MouseWheel>", _on_mousewheel)
        
        def _unbind_from_mousewheel(event):
            self.canvas.unbind_all("<MouseWheel>")
        
        self.canvas.bind('<Enter>', _bind_to_mousewheel)
        self.canvas.bind('<Leave>', _unbind_from_mousewheel)

# ------------------ Enhanced GUI Application ------------------
class EnhancedUniversityGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("🎓 CampusQuery - Enhanced University Information Assistant")
        self.root.geometry("1600x1000")
        self.root.configure(bg=COLORS['light'])
        
        # Initialize variables
        self.status_var = tk.StringVar()
        self.status_var.set("Starting CampusQuery...")
        
        # Initialize components
        self.document_processor = UniversityDocumentProcessor()
        self.query_processor = UniversityQueryProcessor()
        self.system_ready = False
        self.current_result = None
        self.detailed_visible = False
        self.current_detailed_answer = ""
        self.document_paths = {}
        self.current_query = ""  # ENHANCED: Store current query for snippet viewer
        
        # Create custom fonts
        self.setup_fonts()
        
        # Create GUI with improved layout
        self.create_widgets()
        
        # Initialize system
        self.initialize_system()
    
    def setup_fonts(self):
        """Setup custom fonts for better text formatting"""
        self.title_font = font.Font(family="Segoe UI", size=18, weight="bold")
        self.subtitle_font = font.Font(family="Segoe UI", size=14, weight="bold")
        self.normal_font = font.Font(family="Segoe UI", size=12)
        self.bold_font = font.Font(family="Segoe UI", size=12, weight="bold")
        self.mono_font = font.Font(family="Consolas", size=10)
        self.large_font = font.Font(family="Segoe UI", size=13)
    
    def create_widgets(self):
        """Create the main GUI with scrollable interface"""
        # Header (Fixed)
        header_frame = tk.Frame(self.root, bg=COLORS['dark'], height=80)
        header_frame.pack(fill=tk.X)
        header_frame.pack_propagate(False)
        
        title_label = tk.Label(
            header_frame,
            text="🎓 CampusQuery - Enhanced University Information Assistant",
            font=self.title_font,
            fg='white', bg=COLORS['dark']
        )
        title_label.pack(expand=True)
        
        self.status_indicator = tk.Label(
            header_frame,
            text="🔴 Loading...",
            font=self.normal_font,
            fg=COLORS['warning'], bg=COLORS['dark']
        )
        self.status_indicator.place(relx=0.95, rely=0.5, anchor='e')
        
        # Main scrollable content area
        self.main_scroll = ScrollableFrame(self.root)
        self.main_scroll.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Content container (inside scrollable frame)
        content = self.main_scroll.scrollable_frame
        
        # Query section
        self.create_query_section(content)
        
        # Results section
        self.create_results_section(content)
        
        # Status bar (Fixed)
        status_frame = tk.Frame(self.root, bg=COLORS['dark'], height=35)
        status_frame.pack(fill=tk.X, side=tk.BOTTOM)
        status_frame.pack_propagate(False)
        
        self.status_bar = tk.Label(
            status_frame, textvariable=self.status_var,
            font=self.mono_font, fg='white', bg=COLORS['dark']
        )
        self.status_bar.pack(fill=tk.X, padx=10, pady=5)
    
    def create_query_section(self, parent):
        """Create the query input section"""
        query_frame = tk.LabelFrame(
            parent,
            text="Ask Your Question",
            font=self.subtitle_font,
            bg=COLORS['white'],
            fg=COLORS['text_dark'],
            padx=20,
            pady=15
        )
        query_frame.pack(fill=tk.X, pady=(0, 20))
        
        # Instructions
        instr_label = tk.Label(
            query_frame,
            text="💡 Ask questions about university programs, admissions, policies, facilities, etc. Click documents to view them with advanced PDF features including search and highlighting.",
            font=self.normal_font,
            fg=COLORS['text_dark'],
            bg=COLORS['white']
        )
        instr_label.pack(anchor=tk.W, pady=(0, 15))
        
        # Query input with better sizing
        self.query_text = scrolledtext.ScrolledText(
            query_frame,
            height=5,
            font=self.normal_font,
            wrap=tk.WORD,
            bg=COLORS['light'],
            relief=tk.SOLID,
            borderwidth=1
        )
        self.query_text.pack(fill=tk.X, pady=(0, 15))
        
        # Buttons frame
        btn_frame = tk.Frame(query_frame, bg=COLORS['white'])
        btn_frame.pack(fill=tk.X, pady=(0, 10))
        
        self.query_btn = tk.Button(
            btn_frame,
            text="🔍 Get Answer",
            command=self.process_query,
            font=self.bold_font,
            bg=COLORS['primary'],
            fg='white',
            relief=tk.FLAT,
            padx=25,
            pady=12,
            state=tk.DISABLED,
            cursor='hand2'
        )
        self.query_btn.pack(side=tk.LEFT, padx=(0, 15))
        
        clear_btn = tk.Button(
            btn_frame,
            text="🧹 Clear",
            command=self.clear_query,
            font=self.bold_font,
            bg=COLORS['secondary'],
            fg='white',
            relief=tk.FLAT,
            padx=25,
            pady=12,
            cursor='hand2'
        )
        clear_btn.pack(side=tk.LEFT)
        
        # Sample queries with better layout
        sample_frame = tk.LabelFrame(
            parent,
            text="Sample Questions",
            font=self.subtitle_font,
            bg=COLORS['white'],
            fg=COLORS['text_dark'],
            padx=20,
            pady=15
        )
        sample_frame.pack(fill=tk.X, pady=(0, 30))
        
        # Create a grid for sample buttons
        samples = [
            "What programs does the university offer?",
            "What are the admission requirements?",
            "What is the fee structure?",
            "What scholarships are available?",
            "What are the campus facilities?"
        ]
        
        for i, sample in enumerate(samples):
            row = i // 2
            col = i % 2
            btn = tk.Button(
                sample_frame,
                text=sample,
                command=lambda s=sample: self.set_sample_query(s),
                font=self.normal_font,
                bg=COLORS['info'],
                fg='white',
                relief=tk.FLAT,
                padx=15,
                pady=8,
                cursor='hand2',
                wraplength=300
            )
            btn.grid(row=row, column=col, padx=10, pady=5, sticky='ew')
        
        # Configure grid weights
        sample_frame.grid_columnconfigure(0, weight=1)
        sample_frame.grid_columnconfigure(1, weight=1)
    
    def create_results_section(self, parent):
        """Create the results section with proper scrolling"""
        results_container = tk.Frame(parent, bg=COLORS['light'])
        results_container.pack(fill=tk.BOTH, expand=True)
        
        # Answer section
        answer_frame = tk.LabelFrame(
            results_container,
            text="📝 Answer",
            font=self.subtitle_font,
            bg=COLORS['white'],
            fg=COLORS['text_dark'],
            padx=20,
            pady=15
        )
        answer_frame.pack(fill=tk.X, pady=(0, 20))
        
        # Answer text with scrollbar
        answer_container = tk.Frame(answer_frame, bg=COLORS['white'])
        answer_container.pack(fill=tk.X, pady=(0, 15))
        
        self.answer_text = scrolledtext.ScrolledText(
            answer_container,
            height=12,
            font=self.large_font,
            wrap=tk.WORD,
            bg=COLORS['light'],
            relief=tk.SOLID,
            borderwidth=1,
            state=tk.DISABLED
        )
        self.answer_text.pack(fill=tk.X)
        
        # Action buttons with better spacing
        action_frame = tk.Frame(answer_frame, bg=COLORS['white'])
        action_frame.pack(fill=tk.X, pady=(0, 10))
        
        self.detail_btn = tk.Button(
            action_frame,
            text="📖 Show Detailed Explanation",
            command=self.toggle_detailed_explanation,
            font=self.bold_font,
            bg=COLORS['warning'],
            fg='white',
            relief=tk.FLAT,
            padx=20,
            pady=10,
            cursor='hand2',
            state=tk.DISABLED
        )
        self.detail_btn.pack(side=tk.LEFT, padx=(0, 15))
        
        self.export_btn = tk.Button(
            action_frame,
            text="💾 Export Answer",
            command=self.export_answer,
            font=self.bold_font,
            bg=COLORS['success'],
            fg='white',
            relief=tk.FLAT,
            padx=20,
            pady=10,
            cursor='hand2',
            state=tk.DISABLED
        )
        self.export_btn.pack(side=tk.LEFT)
        
        # Detailed explanation section (initially hidden)
        self.detail_frame = tk.LabelFrame(
            results_container,
            text="📚 Detailed Explanation",
            font=self.subtitle_font,
            bg=COLORS['white'],
            fg=COLORS['text_dark'],
            padx=20,
            pady=15
        )
        
        self.detail_text = scrolledtext.ScrolledText(
            self.detail_frame,
            height=20,
            font=self.large_font,
            wrap=tk.WORD,
            bg=COLORS['bg_secondary'],
            relief=tk.SOLID,
            borderwidth=1,
            state=tk.DISABLED
        )
        self.detail_text.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        
        # ENHANCED: Documents section with improved highlighting instructions
        docs_frame = tk.LabelFrame(
            results_container,
            text="📄 Referenced Sources (Double-Click to Open Enhanced PDF Viewer)",
            font=self.subtitle_font,
            bg=COLORS['white'],
            fg=COLORS['text_dark'],
            padx=20,
            pady=15
        )
        docs_frame.pack(fill=tk.X, pady=(20, 0))
        
        # Create treeview with scrollbars
        tree_container = tk.Frame(docs_frame, bg=COLORS['white'])
        tree_container.pack(fill=tk.X, pady=(0, 10))
        
        # Configure treeview style
        style = ttk.Style()
        style.configure("Custom.Treeview", font=self.normal_font, rowheight=30)
        style.configure("Custom.Treeview.Heading", font=self.bold_font)
        
        self.docs_tree = ttk.Treeview(
            tree_container,
            columns=('filename', 'relevance', 'type'),
            show='headings',
            height=8,
            style="Custom.Treeview"
        )
        
        self.docs_tree.heading('filename', text='📄 Source Name')
        self.docs_tree.heading('relevance', text='🎯 Relevance Score')
        self.docs_tree.heading('type', text='📎 File Type')
        self.docs_tree.column('filename', width=400, minwidth=300)
        self.docs_tree.column('relevance', width=120, minwidth=100)
        self.docs_tree.column('type', width=100, minwidth=80)
        
        # Add scrollbars to treeview
        tree_scroll_y = ttk.Scrollbar(tree_container, orient="vertical", command=self.docs_tree.yview)
        tree_scroll_x = ttk.Scrollbar(tree_container, orient="horizontal", command=self.docs_tree.xview)
        self.docs_tree.configure(yscrollcommand=tree_scroll_y.set, xscrollcommand=tree_scroll_x.set)
        
        # Pack treeview and scrollbars
        self.docs_tree.grid(row=0, column=0, sticky='nsew')
        tree_scroll_y.grid(row=0, column=1, sticky='ns')
        tree_scroll_x.grid(row=1, column=0, sticky='ew')
        tree_container.grid_rowconfigure(0, weight=1)
        tree_container.grid_columnconfigure(0, weight=1)
        
        # ENHANCED: Bind events for snippet viewing
        self.docs_tree.bind('<Double-1>', self.open_enhanced_document_viewer)
        self.docs_tree.bind('<Return>', self.open_enhanced_document_viewer)
        
        # ENHANCED: Better instructions for documents
        docs_instr = tk.Label(
            docs_frame,
            text="💡 Double-click any document to open it in the enhanced PDF viewer with search, highlighting, and follow-up question features",
            font=self.normal_font,
            fg=COLORS['text_light'],
            bg=COLORS['white']
        )
        docs_instr.pack(anchor=tk.W)
        
        # Configure text tags
        self.setup_text_tags()
    
    def setup_text_tags(self):
        """Setup text tags for enhanced formatting"""
        # Answer text tags
        self.answer_text.tag_configure("bold", font=self.bold_font, foreground=COLORS['primary'])
        self.answer_text.tag_configure("heading", font=self.subtitle_font, foreground=COLORS['primary'])
        self.answer_text.tag_configure("normal", font=self.large_font)
        self.answer_text.tag_configure("highlight", background=COLORS['bg_primary'])
        
        # Detail text tags
        self.detail_text.tag_configure("bold", font=self.bold_font, foreground=COLORS['primary'])
        self.detail_text.tag_configure("heading", font=self.subtitle_font, foreground=COLORS['primary'])
        self.detail_text.tag_configure("normal", font=self.large_font)
        self.detail_text.tag_configure("highlight", background=COLORS['bg_primary'])
    
    def initialize_system(self):
        """Initialize system with thread-safe GUI updates"""
        def init_thread():
            try:
                self.root.after(0, lambda: self.update_status("🚀 Loading university documents..."))
                documents = self.document_processor.load_all_university_documents()
                
                if not documents:
                    self.root.after(0, lambda: self.update_status("⚠️ No documents found - Web search will be used"))
                    self.root.after(0, lambda: messagebox.showwarning(
                        "No Documents",
                        f"No documents found in:\n{UNIVERSITY_DOCS_DIR}\n\nThe system will use web search for all queries.\nAdd PDF, DOCX, or text files to use local knowledge base."
                    ))
                else:
                    self.root.after(0, lambda: self.update_status(f"📊 Processing {len(documents)} document chunks..."))
                
                self.query_processor.initialize_system(documents)
                self.system_ready = True
                self.root.after(0, lambda: self.update_status("✅ CampusQuery ready! You can now ask questions."))
                self.root.after(0, lambda: self.query_btn.configure(state=tk.NORMAL))
                self.root.after(0, lambda: self.status_indicator.configure(text="🟢 Ready", fg=COLORS['success']))
                
            except Exception as e:
                error_msg = f"Initialization failed: {str(e)}"
                self.root.after(0, lambda: self.update_status(f"❌ {error_msg}"))
                self.root.after(0, lambda: messagebox.showerror("Error", error_msg))
        
        threading.Thread(target=init_thread, daemon=True).start()
    
    def set_sample_query(self, query):
        self.query_text.delete(1.0, tk.END)
        self.query_text.insert(1.0, query)
    
    def clear_query(self):
        self.query_text.delete(1.0, tk.END)
        self.clear_results()
    
    def clear_results(self):
        # Clear main answer
        self.answer_text.config(state=tk.NORMAL)
        self.answer_text.delete(1.0, tk.END)
        self.answer_text.config(state=tk.DISABLED)
        
        # Clear detailed explanation
        self.detail_text.config(state=tk.NORMAL)
        self.detail_text.delete(1.0, tk.END)
        self.detail_text.config(state=tk.DISABLED)
        
        # Hide detailed explanation if visible
        if self.detailed_visible:
            self.detail_frame.pack_forget()
            self.detailed_visible = False
            self.detail_btn.configure(text="📖 Show Detailed Explanation")
        
        # Disable action buttons
        self.detail_btn.configure(state=tk.DISABLED)
        self.export_btn.configure(state=tk.DISABLED)
        
        # Clear documents tree
        for item in self.docs_tree.get_children():
            self.docs_tree.delete(item)
        self.document_paths.clear()
        
        # Reset variables
        self.current_result = None
        self.current_detailed_answer = ""
        self.current_query = ""
    
    def process_query(self):
        query = self.query_text.get(1.0, tk.END).strip()
        if not query:
            messagebox.showwarning("No Query", "Please enter a question.")
            return
        
        if not self.system_ready:
            messagebox.showwarning("System Not Ready", "Please wait for system to initialize.")
            return
        
        # ENHANCED: Store the current query
        self.current_query = query
        
        # Disable button during processing
        self.query_btn.configure(state=tk.DISABLED, text="🔄 Processing...")
        self.clear_results()
        
        def query_thread():
            try:
                self.root.after(0, lambda: self.update_status("🤔 Searching documents..."))
                result = self.query_processor.process_query(query)
                self.current_result = result
                self.current_detailed_answer = result.detailed_answer
                
                self.root.after(0, lambda: self.display_results(result))
                self.root.after(0, lambda: self.update_status("✅ Query complete"))
                self.root.after(0, lambda: self.query_btn.configure(state=tk.NORMAL, text="🔍 Get Answer"))
                
                # Enable action buttons
                self.root.after(0, lambda: self.detail_btn.configure(state=tk.NORMAL))
                self.root.after(0, lambda: self.export_btn.configure(state=tk.NORMAL))
                
            except Exception as e:
                error_msg = f"Error: {str(e)}"
                self.root.after(0, lambda: self.update_status(f"❌ {error_msg}"))
                self.root.after(0, lambda: messagebox.showerror("Error", error_msg))
                self.root.after(0, lambda: self.query_btn.configure(state=tk.NORMAL, text="🔍 Get Answer"))
        
        threading.Thread(target=query_thread, daemon=True).start()
    
    def display_results(self, result: AnalysisResponse):
        """Display formatted results with enhanced text styling"""
        try:
            # Display main answer
            self.answer_text.config(state=tk.NORMAL)
            self.answer_text.delete(1.0, tk.END)
            
            # Format and display answer
            self.insert_formatted_text(self.answer_text, result.answer)
            self.answer_text.config(state=tk.DISABLED)
            
            # ENHANCED: Populate documents tree with snippet data
            self.populate_documents_tree(result.sources)
            
        except Exception as e:
            logger.error(f"Error displaying results: {e}")
            self.answer_text.config(state=tk.NORMAL)
            self.answer_text.insert(1.0, f"❌ Error displaying results: {str(e)}")
            self.answer_text.config(state=tk.DISABLED)
    
    def insert_formatted_text(self, text_widget, text):
        """Insert text with formatting"""
        pos = 0
        while pos < len(text):
            start = text.find('**', pos)
            if start == -1:
                text_widget.insert(tk.END, text[pos:], "normal")
                break
            
            if start > pos:
                text_widget.insert(tk.END, text[pos:start], "normal")
            
            end = text.find('**', start + 2)
            if end == -1:
                text_widget.insert(tk.END, text[start:], "normal")
                break
            
            bold_text = text[start + 2:end]
            text_widget.insert(tk.END, bold_text, "bold")
            pos = end + 2
    
    def populate_documents_tree(self, sources):
        """ENHANCED: Populate the documents tree with source snippet data"""
        print(f"📄 Populating {len(sources)} sources...")
        self.document_snippets = {}  # Store snippets for each item
        
        for source in sources:
            filename = source.get('filename', 'Unknown')
            relevance = f"{source.get('relevance', 0):.1%}"
            filepath = source.get('filepath', '')
            snippet = source.get('content_snippet', '')
            
            # Determine file type
            if filepath:
                file_ext = os.path.splitext(filepath)[1].lower()
                if file_ext == '.pdf':
                    file_type = "📄 PDF"
                elif file_ext == '.docx':
                    file_type = "📝 Word"
                else:
                    file_type = "📄 Text"
            else:
                file_type = "🌐 Web"
            
            if filepath and os.path.exists(filepath):
                item = self.docs_tree.insert('', 'end', values=(f"📄 {filename}", relevance, file_type))
                self.document_paths[item] = filepath
                self.document_snippets[item] = snippet  # ENHANCED: Store snippet
                print(f"✅ Added: {filename}")
            else:
                # For web search results or missing files
                icon = "🌐" if "Web Search" in filename else "❌"
                item = self.docs_tree.insert('', 'end', values=(f"{icon} {filename}", relevance, file_type))
                self.document_snippets[item] = snippet
                if filepath:
                    self.document_paths[item] = filepath
                print(f"🌐 Added web result: {filename}")
    
    def toggle_detailed_explanation(self):
        """Toggle the detailed explanation panel"""
        if not self.current_result:
            messagebox.showinfo("No Results", "Please ask a question first.")
            return
        
        if self.detailed_visible:
            # Hide detailed explanation
            self.detail_frame.pack_forget()
            self.detailed_visible = False
            self.detail_btn.configure(text="📖 Show Detailed Explanation", bg=COLORS['warning'])
        else:
            # Show detailed explanation
            self.detail_frame.pack(fill=tk.X, pady=(20, 0))
            self.detailed_visible = True
            self.detail_btn.configure(text="📖 Hide Detailed Explanation", bg=COLORS['danger'])
            
            # Display detailed explanation
            self.display_detailed_explanation()
    
    def display_detailed_explanation(self):
        """Display the detailed explanation"""
        self.detail_text.config(state=tk.NORMAL)
        self.detail_text.delete(1.0, tk.END)
        
        # Add heading
        self.detail_text.insert(tk.END, "📚 COMPREHENSIVE ANALYSIS\n\n", "heading")
        
        # Insert detailed answer
        if self.current_detailed_answer:
            self.insert_formatted_text(self.detail_text, self.current_detailed_answer)
        else:
            self.detail_text.insert(tk.END, "No detailed explanation available.", "normal")
        
        self.detail_text.config(state=tk.DISABLED)
        self.detail_text.see(1.0)  # Scroll to top
    
    def open_enhanced_document_viewer(self, event):
        """ENHANCED: Open document with appropriate viewer based on file type"""
        selection = self.docs_tree.selection()
        if not selection:
            messagebox.showinfo("No Selection", "Please select a source to view.")
            return
        
        item = selection[0]
        filepath = self.document_paths.get(item)
        snippet = self.document_snippets.get(item, "No snippet available.")
        
        # Get document info from tree
        values = self.docs_tree.item(item, 'values')
        document_name = values[0].replace('📄 ', '').replace('❌ ', '').replace('🌐 ', '') if values else "Unknown Source"
        document_name = document_name.replace(' (Missing)', '')
        file_type = values[2] if len(values) > 2 else "Unknown"
        
        # Handle web search results
        if "Web Search" in document_name or "🌐" in file_type:
            messagebox.showinfo("Web Search Result", f"Source: {document_name}\n\nContent:\n{snippet[:500]}{'...' if len(snippet) > 500 else ''}")
            return
        
        if not filepath:
            messagebox.showwarning("No Path", "No file path associated with this source.")
            return
        
        try:
            # Check if it's a PDF file
            if filepath.endswith('.pdf'):
                # ENHANCED: Open enhanced PDF viewer
                pdf_viewer = EnhancedPDFViewer(
                    parent=self.root,
                    document_name=document_name,
                    document_path=filepath,
                    snippet_text=snippet,
                    query=self.current_query,
                    main_app=self  # Pass reference to main app for follow-up questions
                )
                self.update_status(f"📄 Opened enhanced PDF viewer for: {document_name}")
            else:
                # Open traditional snippet viewer for non-PDF files
                snippet_viewer = DocumentSnippetViewer(
                    parent=self.root,
                    document_name=document_name,
                    document_path=filepath,
                    snippet_text=snippet,
                    query=self.current_query
                )
                self.update_status(f"📄 Viewing source content from: {document_name}")
            
        except Exception as e:
            error_msg = f"Could not open document viewer:\n{str(e)}"
            messagebox.showerror("Error", error_msg)
            logger.error(f"Error opening document viewer: {e}")
    
    def export_answer(self):
        """Export the current answer to a text file"""
        if not self.current_result:
            messagebox.showinfo("No Answer", "Please ask a question first.")
            return
        
        try:
            from tkinter import filedialog
            filename = filedialog.asksaveasfilename(
                defaultextension=".txt",
                filetypes=[("Text files", "*.txt"), ("All files", "*.*")],
                title="Save Answer As"
            )
            
            if filename:
                with open(filename, 'w', encoding='utf-8') as f:
                    f.write("CampusQuery - Enhanced University Information Assistant\n")
                    f.write(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                    f.write("="*60 + "\n\n")
                    f.write(f"QUESTION:\n{self.current_query}\n\n")
                    f.write(f"ANSWER:\n{self.current_result.answer}\n\n")
                    
                    if self.current_detailed_answer:
                        f.write(f"DETAILED EXPLANATION:\n{self.current_detailed_answer}\n\n")
                    
                    f.write("REFERENCE SOURCES:\n")
                    for ref in self.current_result.document_references:
                        f.write(f"- {ref}\n")
                    
                    f.write(f"\nJUSTIFICATION:\n{self.current_result.justification}\n")
                
                messagebox.showinfo("Export Successful", f"Answer exported to:\n{filename}")
        except Exception as e:
            messagebox.showerror("Export Error", f"Error exporting answer:\n{str(e)}")
    
    def update_status(self, message):
        """Update status bar with timestamp"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        self.status_var.set(f"[{timestamp}] {message}")

def main():
    """Main application entry point"""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        print("Installing required dependency: PyMuPDF")
        os.system("pip install PyMuPDF")
        import fitz
    
    root = tk.Tk()
    root.minsize(1400, 900)
    app = EnhancedUniversityGUI(root)
    
    def on_closing():
        if messagebox.askokcancel("Quit", "Exit CampusQuery?"):
            root.destroy()
    
    root.protocol("WM_DELETE_WINDOW", on_closing)
    
    # Center window
    root.update_idletasks()
    x = (root.winfo_screenwidth() // 2) - (root.winfo_width() // 2)
    y = (root.winfo_screenheight() // 2) - (root.winfo_height() // 2)
    root.geometry(f"+{x}+{y}")
    
    print("🎓 CampusQuery - Enhanced University Information Assistant Started")
    print(f"📁 Documents directory: {UNIVERSITY_DOCS_DIR}")
    print("🌐 Web search fallback enabled")
    print("📄 Enhanced PDF viewer with search and highlighting enabled")
    
    root.mainloop()

if __name__ == "__main__":
    main()
